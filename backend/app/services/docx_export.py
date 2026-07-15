"""Build a 細明體 (PMingLiU) DOCX diagnostic report from sample data.

Layout follows docs/reference (報告模板20260506.docx) — five top-level
sections (一/二/三/四/五) with sub-blocks for SNV/Indel, Mitochondrial,
and CNV variants. Reviewer state (status=1 → Causative, status=2 →
Other; status=C/Candidate is NOT included in the formal report) drives
which variants appear; gene-list mode controls how the "本次檢測基因
包括" section in §五 is rendered.

The font is forced to PMingLiU on every run + on Normal style; Word
will substitute when opening on a machine without 細明體 installed.
"""
from __future__ import annotations

import io
import html
import re
import unicodedata
from typing import Iterable

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from ..config import TERTIARY_OUTPUT_ROOT  # noqa: F401 (kept for callers)
from . import (
    cnv_sv_merge,
    hpo_ontology,
    panel_deadzone,
    phenotype_scorer,
    ploidy,
    report_store,
    sample_loader,
)

# 細明體 (MingLiU) — the *monospace* CJK family. (新細明體 = PMingLiU
# is proportional, which would break the ASCII tables below.) Word
# falls back to whatever the system has if MingLiU isn't installed.
REPORT_FONT     = "MingLiU"
TITLE_FONT_SIZE = Pt(13)
BODY_FONT_SIZE  = Pt(11)

# ClinVar release these reports reference. Hardcoded — bump when the
# clinvar.vcf.gz used by annotate_clinvar.py is refreshed. Format
# matches the template: ISO-like "YYYYMMDD" and natural "YYYY 年 MM 月
# DD 日".
CLINVAR_DATE      = "20260510"
CLINVAR_DATE_HUMN = "2026 年 05 月 10 日"


# ── Lookup tables ─────────────────────────────────────────────────

# VEP SO term → glossary key. Severity order is the standard VEP
# severity ranking; we pick the first match among the variant's
# `Consequence` field (which VEP joins with "&" or ",").
_CONSEQUENCE_SEVERITY: list[tuple[str, str]] = [
    ("stop_gained",                       "nonsense"),
    ("frameshift_variant",                "frameshift"),
    ("splice_acceptor_variant",           "splice"),
    ("splice_donor_variant",              "splice"),
    ("splice_region_variant",             "splice"),
    ("start_lost",                        "start_loss"),
    ("stop_lost",                         "stop_loss"),
    ("inframe_insertion",                 "in_frame"),
    ("inframe_deletion",                  "in_frame"),
    ("missense_variant",                  "missense"),
    ("missense",                          "missense"),    # mito script short form
    ("synonymous_variant",                "synonymous"),
    ("synonymous",                        "synonymous"),  # mito script short form
    ("5_prime_UTR_variant",               "noncoding"),
    ("3_prime_UTR_variant",               "noncoding"),
    ("intron_variant",                    "noncoding"),
    ("non_coding_transcript_exon_variant","noncoding"),
    ("non_coding_transcript_variant",     "noncoding"),
    ("upstream_gene_variant",             "noncoding"),
    ("downstream_gene_variant",           "noncoding"),
    ("regulatory_region_variant",         "noncoding"),
    ("intergenic_variant",                "noncoding"),
]

# Glossary key → 中文制式句 (from 報告模板20260506.docx lines 81-89)
_CONSEQUENCE_ZH: dict[str, str] = {
    "nonsense":   "無義突變 (Nonsense mutation)，此變異會形成過早的終止密碼子，使蛋白質轉譯提前終止，通常導致蛋白質功能喪失。",
    "missense":   "誤義突變 (Missense mutation)，此變異使密碼子改變，造成氨基酸改變，可能影響蛋白質的結構穩定性、功能或相互作用。",
    "synonymous": "同義突變 (Synonymous mutation)，此變異雖不改變蛋白質胺基酸序列，但可能影響 mRNA 剪接效率、穩定性或轉譯效率，因此仍可能具有功能影響。",
    "frameshift": "移碼突變 (Frameshift mutation)，此變異會改變原有閱讀框架，造成後續胺基酸順序大幅改變，通常導致不完整或功能喪失的蛋白質。",
    "in_frame":   "非移碼突變 (In-frame mutation)，此變異不改變閱讀框架，但會增加或刪除完整密碼子，可能影響蛋白質結構、穩定性或功能性區域。",
    "splice":     "剪接位點突變 (Splice site mutation)，此變異可能破壞正常 RNA 剪接，造成外顯子缺失、內含子保留或異常剪接，進而改變蛋白質長度或功能。",
    "start_loss": "起始密碼子缺失 (Start-loss mutation)，此變異可能影響正常翻譯啟動，使蛋白質無法正常生成，或自替代啟動位置開始產生異常蛋白。",
    "stop_loss":  "終止密碼子缺失 (Stop-loss mutation)，此變異會破壞終止密碼子，使轉譯延伸至下游序列生成額外胺基酸，可能改變蛋白質長度與功能。",
    "noncoding":  "非編碼區變異 (Noncoding region mutation)，此變異不直接改變蛋白序列，但可能影響基因調控、剪接、mRNA 穩定性或表達量。",
}

# OMIM inheritance code → 中文
_INHERITANCE_ZH: dict[str, str] = {
    "AD":   "體染色體顯性遺傳",
    "AR":   "體染色體隱性遺傳",
    "XL":   "性聯遺傳",
    "XLD":  "性染色體顯性遺傳",
    "XLR":  "性染色體隱性遺傳",
    "YL":   "Y 染色體連鎖遺傳",
    "MT":   "粒線體遺傳",
    "Mi":   "粒線體遺傳",
    "Mu":   "多因子遺傳",
    "Isol": "散發性",
}

_INHERITANCE_SUPPRESS = {"DR", "DD", "Smu", "SMU", "Digenic", "Somatic"}

_ZYGOSITY_ZH: dict[str, str] = {
    "Heterozygous":    "異合子",
    "Homozygous":      "同合子",
    "Hemizygous":      "半合子",
    "het":             "異合子",
    "hom":             "同合子",
    "Heteroplasmic":   "異質性",
    "Homoplasmic":     "同質性",
}

_CNV_KIND_ZH: dict[str, str] = {
    "DEL": "缺失",
    "DUP": "重複",
    "INS": "插入",
    "INV": "倒轉",
}

# Amino-acid single-letter → three-letter mapping. Used to normalise
# MITOMAP's compact `A64V` style to HGVS `p.Ala64Val` so the report
# matches the SNV nucleotide column format.
_AA_1_TO_3: dict[str, str] = {
    "A": "Ala", "R": "Arg", "N": "Asn", "D": "Asp", "C": "Cys",
    "E": "Glu", "Q": "Gln", "G": "Gly", "H": "His", "I": "Ile",
    "L": "Leu", "K": "Lys", "M": "Met", "F": "Phe", "P": "Pro",
    "S": "Ser", "T": "Thr", "W": "Trp", "Y": "Tyr", "V": "Val",
    "*": "Ter", "X": "Ter", "U": "Sec", "O": "Pyl",
}

# AnnotSV's ACMG/ClinGen classification is an int 1-5 on full rows.
_CNV_ACMG_INT_TO_LABEL: dict[int, str] = {
    1: "Benign",
    2: "Likely benign",
    3: "Uncertain significance",
    4: "Likely pathogenic",
    5: "Pathogenic",
}


def _consequence_zh(vep_csq: str) -> str:
    """Pick the most severe glossary entry for a VEP Consequence string.

    The Consequence field is `&`- or `,`-joined VEP SO terms. We walk
    `_CONSEQUENCE_SEVERITY` (already in severity order) and return the
    first 中文 句 whose term appears in the input.
    """
    if not vep_csq:
        return _CONSEQUENCE_ZH["noncoding"]
    tokens = set(re.split(r"[&,;|/\s]+", vep_csq))
    for so_term, key in _CONSEQUENCE_SEVERITY:
        if so_term in tokens:
            return _CONSEQUENCE_ZH[key]
    return _CONSEQUENCE_ZH["noncoding"]


def _inheritance_zh(code: str) -> str:
    """OMIM inheritance code → 中文. Unknown codes pass through verbatim
    so the reviewer can spot weird values (rather than silently mapping
    to a wrong term).
    """
    if not code:
        return ""
    # OMIM inheritance is sometimes a multi-code "AD; AR" — split and
    # translate each, join with 「、」.
    out = []
    for part in re.split(r"[;,/\s]+", code.strip()):
        if part in _INHERITANCE_SUPPRESS:
            continue
        out.append(_INHERITANCE_ZH.get(part, part))
    return "、".join(p for p in out if p)


def _zygosity_zh(z: str) -> str:
    return _ZYGOSITY_ZH.get((z or "").strip(), z or "")


_ZYG_LONG: dict[str, str] = {
    "het":          "Heterozygous",
    "hom":          "Homozygous",
    "hemi":         "Hemizygous",
    "homo":         "Homozygous",
    "heterozygous": "Heterozygous",
    "homozygous":   "Homozygous",
    "hemizygous":   "Hemizygous",
}


def _zygosity_long(z: str) -> str:
    """Expand short forms (het/hom/hemi) to the canonical English long
    form used in the diagnostic report tables."""
    s = (z or "").strip()
    return _ZYG_LONG.get(s.lower(), s)


def _strip_tx_prefix(s: str) -> str:
    """Pipeline HGVS values sometimes come prefixed with the transcript
    ID (e.g. "NM_000295.5:c.1075A>G"); the report's 核苷酸 column shows
    transcript-only on the gene header line, so strip the prefix here.
    """
    return re.sub(r"^[A-Z]+_\d+(\.\d+)?:", "", s or "")


def _mito_aa_to_hgvsp(aa: str) -> str:
    """MITOMAP gives "A64V" — convert to "p.Ala64Val" so the report's
    nucleotide cell matches the SNV HGVS format. Already-3-letter input
    (with or without `p.` prefix) is normalised to carry the `p.`.
    Empty / unparseable input passes through.
    """
    if not aa:
        return ""
    s = aa.strip()
    if s.startswith("p."):
        s = s[2:]
    # 1-letter form: A64V or *64L (stop)
    m1 = re.match(r"^([A-Z*])(\d+)([A-Z*])$", s)
    if m1:
        a1, pos, a2 = m1.groups()
        return f"p.{_AA_1_TO_3.get(a1, a1)}{pos}{_AA_1_TO_3.get(a2, a2)}"
    # 3-letter form: Ala64Val
    m3 = re.match(r"^([A-Z][a-z]{2}|Ter)(\d+)([A-Z][a-z]{2}|Ter)$", s)
    if m3:
        return f"p.{s}"
    return aa


def _structure_label(v: dict) -> str:
    """VEP's EXON / INTRON fields look like "5/22" (rank/total). The
    report writes them as `exon5` or `intron6` — pick whichever has
    a value and prefix with the type."""
    def rank(value) -> str:
        text = str(value or "").strip()
        return "" if text.upper() in {"", ".", "-", "NA", "N/A"} else text

    exon = rank(v.get("exon"))
    intron = rank(v.get("intron"))
    if exon:
        return "exon" + exon.split("/")[0].strip()
    if intron:
        return "intron" + intron.split("/")[0].strip()
    return ""


# ── Font + paragraph helpers ──────────────────────────────────────

def _set_run_font(run, name: str = REPORT_FONT) -> None:
    """Set both Western and East-Asian fonts on a run."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"),    name)
    rFonts.set(qn("w:hAnsi"),    name)


def _add_paragraph(doc, text: str, bold: bool = False,
                   size: Pt = BODY_FONT_SIZE, align: str | None = None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = 1
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    _set_run_font(run)
    return p


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append clickable display text for an external URL."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_properties.append(run_style)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _blank(doc):
    p = doc.add_paragraph()
    run = p.add_run("")
    run.font.size = BODY_FONT_SIZE
    _set_run_font(run)


def _apply_normal_font(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = REPORT_FONT
    normal.font.size = BODY_FONT_SIZE
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), REPORT_FONT)
    rFonts.set(qn("w:ascii"),    REPORT_FONT)
    rFonts.set(qn("w:hAnsi"),    REPORT_FONT)
    # Word's default Normal style adds 8pt space after each paragraph;
    # the reviewer was manually clearing that on every export. Set it
    # to 0 so the report opens already tight.
    pf = normal.paragraph_format
    pf.space_after  = Pt(0)
    pf.space_before = Pt(0)
    # Tighten line spacing too — single (1.0) instead of 1.15 multiple.
    from docx.shared import Pt as _Pt  # noqa: F401 (silence linter)
    pf.line_spacing = 1.0


def _apply_page_margins(doc) -> None:
    """A4 portrait with 1.5cm L/R margins so the 89-char-wide variant
    tables (4-space indent + 85-char ===== box) fit on one line in
    細明體 11pt."""
    for section in doc.sections:
        section.left_margin  = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)


# ── ASCII-table helpers ───────────────────────────────────────────
# 細明體 renders CJK chars at 2x ASCII width and ASCII chars at uniform
# half-width, so a fixed-width column layout drawn with spaces lines up
# visually. We treat East-Asian Wide/Fullwidth (W/F) as 2 columns and
# everything else as 1.

def _ea_width(c: str) -> int:
    return 2 if unicodedata.east_asian_width(c) in ("W", "F") else 1


def _str_width(s: str) -> int:
    return sum(_ea_width(c) for c in s)


def _pad_right(s: str, w: int) -> str:
    """Right-pad with spaces up to display width w."""
    return s + " " * max(0, w - _str_width(s))


def _wrap_to_cols(text: str, width: int, mode: str = "char") -> list[str]:
    """Wrap `text` so each chunk fits within `width` display columns.

    mode="char" → greedy char-by-char wrap respecting CJK double width
                  (default; mid-token break, matches template's HGVS
                  wrap behaviour).
    mode="token" → one whitespace-delimited token per line; also splits
                   after "/" so "Pathogenic/Likely pathogenic" becomes
                   3 lines (Pathogenic/ · Likely · pathogenic).
                   Tokens longer than `width` fall back to char wrap.
    mode="word" → greedy word wrap for prose cells, preserving words
                  whenever possible.
    mode="buffered" → greedy char wrap with one trailing display column
                      reserved as whitespace before the next table cell.
    mode="token-buffered" → token wrap with one trailing display column
                            reserved before the next table cell.
    """
    if text is None:
        return [""]
    s = str(text)
    if not s:
        return [""]

    if mode == "token":
        # Split on whitespace; also break after "/" so combined sigs
        # like "Pathogenic/Likely pathogenic" land on three lines.
        tokens = [t for t in re.split(r"\s+|(?<=/)", s) if t]
        out: list[str] = []
        for tok in tokens:
            if _str_width(tok) <= width:
                out.append(tok)
            else:
                out.extend(_wrap_to_cols(tok, width, mode="char"))
        return out or [""]

    if mode == "token-buffered":
        return _wrap_to_cols(s, max(1, width - 1), mode="token")

    if mode == "word":
        tokens = [t for t in re.split(r"\s+", s) if t]
        out: list[str] = []
        cur = ""
        cur_w = 0
        for tok in tokens:
            tok_w = _str_width(tok)
            if tok_w > width:
                if cur:
                    out.append(cur)
                    cur = ""
                    cur_w = 0
                out.extend(_wrap_to_cols(tok, width, mode="char"))
                continue
            sep = " " if cur else ""
            sep_w = 1 if cur else 0
            if cur and cur_w + sep_w + tok_w > width:
                out.append(cur)
                cur = tok
                cur_w = tok_w
            else:
                cur = f"{cur}{sep}{tok}" if cur else tok
                cur_w += sep_w + tok_w
        if cur:
            out.append(cur)
        return out or [""]

    if mode == "word-buffered":
        return _wrap_to_cols(s, max(1, width - 1), mode="word")

    if mode == "genotype-buffered":
        content_width = max(1, width - 1)
        if _str_width(s) > content_width and "/" in s:
            first, rest = s.split("/", 1)
            left = f"{first}/"
            if _str_width(left) <= content_width:
                return [left] + _wrap_to_cols(rest, content_width, mode="char")
        return _wrap_to_cols(s, content_width, mode="char")

    if mode == "hgvs":
        # Keep one blank display column between HGVS content and the next
        # table cell. Protein notation always starts on its own line.
        content_width = max(1, width - 1)
        m = re.search(r"(\(p\.[^)]*\)?)$", s)
        if m and m.start() > 0:
            prefix = s[: m.start()]
            suffix = m.group(1)
            return (
                _wrap_to_cols(prefix, content_width, mode="char")
                + _wrap_to_cols(suffix, content_width, mode="char")
            )
        return _wrap_to_cols(s, content_width, mode="char")

    if mode == "buffered":
        return _wrap_to_cols(s, max(1, width - 1), mode="char")

    out, cur, cur_w = [], "", 0
    for c in s:
        cw = _ea_width(c)
        if cur and cur_w + cw > width:
            out.append(cur)
            cur, cur_w = c, cw
        else:
            cur += c
            cur_w += cw
    if cur:
        out.append(cur)
    return out


def _ascii_table(doc,
                 columns: list[tuple],  # (header, width) or (header, width, mode)
                 rows: list[list[str]],
                 indent: str = "    ") -> None:
    """Render a ==== bounded table where each column has a fixed display
    width; over-long cells wrap per the column's `mode` ("char" default
    / "token" for clinical sigs). Columns are packed without inter-
    column whitespace — pad each column to its full width on its own;
    the data row gets a single space of side-padding inside the ====
    boundary.

    With column widths summing to N and (k) columns, the data line is
    N+2 chars wide (1-space pad each side) and the sep is `=`*(N+2).
    """
    # Normalize each column spec to (header, width, mode).
    cols = [(c[0], c[1], (c[2] if len(c) > 2 else "char")) for c in columns]
    widths = [w for _, w, _ in cols]
    modes  = [m for _, _, m in cols]
    total  = sum(widths) + 2   # +1 pad on each side
    sep    = "=" * total

    def row_lines(cells: list[str], header: bool = False) -> list[str]:
        # Headers don't wrap (kept short by design).
        if header:
            parts = [_pad_right(str(c or ""), w) for c, w in zip(cells, widths)]
            return [f"{indent} {''.join(parts)} "]
        wrapped = [
            _wrap_to_cols(str(c or ""), w, mode=m)
            for c, w, m in zip(cells, widths, modes)
        ]
        n = max((len(w) for w in wrapped), default=1)
        lines = []
        for i in range(n):
            parts = [
                _pad_right(cell_lines[i] if i < len(cell_lines) else "", w)
                for cell_lines, w in zip(wrapped, widths)
            ]
            lines.append(f"{indent} {''.join(parts)} ")
        return lines

    header_lines = [
        f"{indent}{sep}",
        *row_lines([h for h, _, _ in cols], header=True),
        f"{indent}{sep}",
    ]
    _add_paragraph(doc, "\n".join(header_lines))
    for row_index, row in enumerate(rows):
        lines = row_lines(row)
        if row_index == len(rows) - 1:
            lines.append(f"{indent}{sep}")
        _add_paragraph(doc, "\n".join(lines))


# ── Variant-bucketing helpers ─────────────────────────────────────

def _buckets_for_type(variants: dict, statuses: dict, status: str) -> list[dict]:
    """Pull variants whose reviewer-set status matches; preserve order
    from the variants dict (tier-sorted by the adapter).
    """
    out = []
    for vid in variants:
        s = (statuses.get(vid) or "").strip()
        if s == status:
            out.append(variants[vid])
    return out


def _manual_for(report: dict, status: str) -> list[dict]:
    out = []
    for m in (report.get("manual_variants") or []):
        if (m or {}).get("status") == status and (m.get("position") or "").strip():
            out.append(m)
    return out


# ── Sections ──────────────────────────────────────────────────────

def _section_test_info(doc, test_type: str, *, health: bool = False) -> None:
    """一、檢驗項目"""
    is_wgs = (test_type or "").upper() == "WGS"
    if is_wgs:
        label = "次世代定序全基因體定序檢測"
    elif health:
        label = "次世代定序全外顯子定序檢測"
    else:
        label = "次世代定序全外顯子定序檢測-單基因遺傳疾病"
    _add_paragraph(doc, f"一、檢驗項目: {label}")
    _blank(doc)


def _section_panel_set(doc) -> None:
    """二、檢驗套組  (固定「非特定」for now — reviewer 套組概念之後再加)"""
    _add_paragraph(doc, "二、檢驗套組: 非特定")
    _blank(doc)


def _summary_line(name: str, hpo_labels: list[str]) -> str:
    """『在非特定 (term, term, ...) 檢驗套組中未找到已知致病性位點。』"""
    if hpo_labels:
        return f"    在{name} ({', '.join(hpo_labels)}) 檢驗套組中未找到已知致病性位點。"
    return f"    在{name}檢驗套組中未找到已知致病性位點。"


def _section_results(doc, sample: dict, report: dict, test_type: str) -> None:
    """三、檢測結果 — the meat of the report.

    第一類 (status=1) and 第二類 (status=2) sections each contain every
    variant of that status with the matching per-type block (SNV /
    Mito / CNV+SV) rendered inline. No "## SNV/indel" labels — those
    were template structure markers, not actual report headings.
    """
    _add_paragraph(doc, "三、檢測結果")
    _add_paragraph(doc, "  檢體說明:")
    _add_paragraph(doc, "    檢體類別：血液")
    _add_paragraph(doc, "  綜合說明:")

    statuses = report.get("status", {}) or {}
    edits    = report.get("edits", {})  or {}
    is_wgs   = (test_type or "").upper() == "WGS"

    snv_vars  = sample.get("variants", {})       or {}
    cnv_vars, sv_vars = cnv_sv_merge.apply_confirmed_merges(
        sample.get("cnv_variants", {}) or {},
        sample.get("sv_variants", {}) or {},
        report.get("cnv_sv_merges") or [],
    )
    mito_vars = sample.get("mito_variants", {})  or {}

    # Group by reviewer status. Each entry is ("kind", variant_dict).
    # Insertion order = (snv → mito → cnv → sv), so 第一類/第二類 list
    # SNVs first (the most common case in past reports). CNV/SV entries
    # are sorted inside their source by combined phenotype + AnnotSV score.
    def _ranked_items(src: dict, kind: str):
        items = list(src.items())
        if kind in {"cnv", "sv"}:
            items.sort(key=lambda item: (
                -float(item[1].get("cnv_sv_sort_score") or -999),
                str(item[0]),
            ))
        return items

    def _collect(status: str) -> list[tuple[str, dict]]:
        out: list[tuple[str, dict]] = []
        for src_kind, src in (("snv",  snv_vars),
                              ("mito", mito_vars),
                              ("cnv",  cnv_vars),
                              ("sv",   sv_vars)):
            for vid, v in _ranked_items(src, src_kind):
                if (statuses.get(vid) or "").strip() == status:
                    out.append((src_kind, v))
        return out

    bucket1 = _collect("1")
    bucket2 = _collect("2")
    man1 = _manual_for(report, "1")
    man2 = _manual_for(report, "2")

    # — 第一類
    _add_paragraph(doc, "    第一類：與臨床症狀相關基因之已知致病性變異位點")
    if bucket1 or man1:
        for kind, v in bucket1:
            _render_variant(doc, kind, v, tier="1",
                            edits=edits.get(v.get("id", ""), {}),
                            is_wgs=is_wgs)
        for m in man1:
            _render_manual_variant(doc, m)
    else:
        # Empty bucket — match the template's wording. We mirror past
        # reports that surface the HPO labels in the empty notice.
        hpo_labels = [r.get("label", "") or r.get("phenotype", "")
                      for r in (sample.get("patient_phenotype") or [])]
        hpo_labels = [x for x in hpo_labels if x][:8]
        _add_paragraph(doc, _summary_line("非特定", hpo_labels))
    _blank(doc)

    # — 第二類
    _add_paragraph(doc, "    第二類：其他變異位點")
    if bucket2 or man2:
        for kind, v in bucket2:
            _render_variant(doc, kind, v, tier="2",
                            edits=edits.get(v.get("id", ""), {}),
                            is_wgs=is_wgs)
        for m in man2:
            _render_manual_variant(doc, m)
    else:
        _add_paragraph(doc, "    未找到其他變異位點。")
    _blank(doc)

    # Footer recommendation (always shown)
    _add_paragraph(doc, "    建議比對臨床表徵並進行父母親與家族成員之變異位點檢測，"
                        "以釐清上述變異致病之可能性；根據家族成員變異位點檢測報告或"
                        "相關資料庫更新，可能影響變異位點ACMG判讀結果。")
    _blank(doc)
    referenced = bucket1 + bucket2
    if referenced:
        _add_paragraph(doc, "  參考資料:")
        for kind, v in referenced:
            _add_paragraph(doc, _variant_reference_text(
                kind, v, edits=edits.get(v.get("id", ""), {}), is_wgs=is_wgs,
            ))
            _blank(doc)


def _render_variant(doc, kind: str, v: dict, *, tier: str, edits: dict,
                    is_wgs: bool) -> None:
    if kind == "snv":
        _snv_variant_block(doc, v, tier=tier, edits=edits)
    elif kind == "mito":
        _mito_variant_block(doc, v, tier=tier, edits=edits)
    else:  # cnv / sv share the same template
        _cnv_variant_block(doc, v, tier=tier, is_wgs=is_wgs, edits=edits)
    _blank(doc)


def _variant_reference_text(kind: str, v: dict, *, edits: dict,
                            is_wgs: bool) -> str:
    if kind == "snv":
        return _snv_reference_text(v, edits)
    if kind == "mito":
        return _mito_reference_text(v, edits)
    omim_genes = _omim_genes(v)
    report_genes = edits.get("report_genes") or {}
    if isinstance(report_genes, dict):
        kept = [g for g in omim_genes if report_genes.get(g.get("gene"))]
        if kept:
            omim_genes = kept
    return _cnv_reference_text(v, edits, omim_genes, _sv_kind_zh(v), is_wgs)


def _render_manual_variant(doc, m: dict) -> None:
    _add_paragraph(doc, f"    {m.get('position', '')}", bold=True)
    if m.get("disease"):
        _add_paragraph(doc, f"    {m['disease']}")
    if m.get("comment"):
        _add_paragraph(doc, f"    {m['comment']}")
    _blank(doc)


# ── Subsections: per variant type ─────────────────────────────────

def _afs_str(variant: dict) -> str:
    """Return e.g. '0.01%' or '未報導過發生率'. Prefer the most-relevant
    AF column (gnomAD genome > exome).
    """
    for key in ("AF", "AF_exome"):
        v = variant.get(key)
        if v in (None, "", "."):
            continue
        try:
            f = float(v)
            return f"{f * 100:.4f}%".rstrip("0").rstrip(".")
        except ValueError:
            continue
    return ""


def _clinvar_label(variant: dict) -> str:
    sig = (variant.get("CLNSIG") or "").strip()
    return sig.replace("_", " ") if sig and sig != "." else "—"


def _acmg_label(variant: dict, edits_for_v: dict) -> str:
    """Return the ACMG/AMP classification label for any variant kind.

    Source priority (most-specific first):
      1. Mito reviewer manual entry   (edits.ACMG_classification_mito)
      2. SNV reviewer override        (edits.ACMG_classification)
      3. CNV/SV reviewer override     (edits.ACMG_class_sv → int 1-5)
      4. SNV GeneBe value             (variant.genebe_acmg_class)
      5. SNV pipeline value           (variant.ACMG_classification)
      6. CNV/SV pipeline value, int   (variant.acmg_class 1-5)
    Mito has no automated ACMG source — when no manual entry is set
    the function returns an empty string so the report's table cell
    renders blank, per spec.
    """
    cls = (edits_for_v.get("ACMG_classification_mito") or "").strip()
    if cls:
        return cls
    cls = (edits_for_v.get("ACMG_classification") or "").strip()
    if cls:
        return cls
    raw_cnv = (edits_for_v.get("ACMG_class_sv") or "").strip()
    if raw_cnv:
        try:
            n = int(raw_cnv)
            if n in _CNV_ACMG_INT_TO_LABEL:
                return _CNV_ACMG_INT_TO_LABEL[n]
        except ValueError:
            return raw_cnv
    cls = (variant.get("genebe_acmg_class") or "").strip()
    if cls:
        return cls
    cls = (variant.get("ACMG_classification") or "").strip()
    if cls:
        return cls
    n = variant.get("acmg_class")
    if isinstance(n, int) and n in _CNV_ACMG_INT_TO_LABEL:
        return _CNV_ACMG_INT_TO_LABEL[n]
    return ""


def _patho_sentence(acmg_class: str) -> str:
    """Note-2 sentence in §三 driven by the actual ACMG class — so a
    VUS variant reads "此為不確定意義之變異位點..." not "此為致病性...".
    """
    cls = (acmg_class or "").strip().lower().replace("_", " ")
    if cls == "pathogenic":
        return "此為致病性之變異位點，與臨床症狀相關。"
    if cls in ("likely pathogenic",):
        return "此為疑似致病性之變異位點，與臨床症狀相關。"
    if "pathogenic/likely" in cls or "likely pathogenic/pathogenic" in cls:
        return "此為致病性 / 疑似致病性之變異位點，與臨床症狀相關。"
    if cls in ("uncertain significance", "vus"):
        return ("此為不確定意義之變異位點，其臨床意義須由醫師配合其他"
                "相關資料進行最佳綜合判斷。")
    if cls in ("likely benign",):
        return "此為可能良性之變異位點。"
    if cls == "benign":
        return "此為良性之變異位點。"
    # Mito MITOMAP statuses
    if any(k in cls for k in ("cfrm", "confirmed", "[p]", "[lp]")):
        return "此為致病性之變異位點，與臨床症狀相關。"
    if "reported" in cls:
        return ("此為與疾病相關之變異位點，其臨床意義須由醫師配合其他"
                "相關資料進行最佳綜合判斷。")
    # Unknown / empty
    return "此變異位點之臨床意義須由醫師配合其他相關資料進行最佳綜合判斷。"


def _picked_disease_for_snv(v: dict, edits: dict) -> str:
    """Return the reviewer-picked Disease cell, falling back to the first."""
    picked = edits.get("report_diseases") or {}
    if isinstance(picked, dict):
        for idx in range(1, 6):
            if picked.get(str(idx)) or picked.get(idx):
                disease = (v.get(f"Disease{idx}") or "").strip()
                if disease and disease != "NA":
                    return disease
    for idx in range(1, 6):
        disease = (v.get(f"Disease{idx}") or "").strip()
        if disease and disease != "NA":
            return disease
    return (v.get("OMIM_disease") or "").strip()


def _disease_info(disease: str) -> tuple[str, str, str]:
    """Parse disease name, disease-specific inheritance, and phenotype MIM."""
    lines = (disease or "").splitlines()
    first_line = lines[0].strip() if lines else ""
    inheritance_match = re.search(
        r"\((AD|AR|XLD|XLR|XL|YL|MT|Mi|DR|DD|Smu|Mu|Isol)"
        r"(?:\s*[/,;]\s*(?:AD|AR|XLD|XLR|XL|YL|MT|Mi|DR|DD|Smu|Mu|Isol))*\)",
        first_line,
        re.IGNORECASE,
    )
    phenotype_mim_match = re.search(r"\((\d{6})\)", first_line)
    inheritance = inheritance_match.group(0)[1:-1] if inheritance_match else ""
    phenotype_mim = phenotype_mim_match.group(1) if phenotype_mim_match else ""
    metadata_starts = [
        match.start() for match in (phenotype_mim_match, inheritance_match)
        if match
    ]
    name = first_line[:min(metadata_starts)] if metadata_starts else first_line
    return name.rstrip(" :,;").strip(), inheritance, phenotype_mim


def _omim_block_for_snv(v: dict, edits: dict) -> str:
    """『GENE 為 DISEASE 的致病基因之一，其遺傳模式屬於 X
    (Phenotype MIM number: M)』 — fall back to whatever is present."""
    gene = v.get("gene_symbol") or v.get("GENE") or "?"
    disease_str, inheritance, phenotype_mim = _disease_info(
        _picked_disease_for_snv(v, edits)
    )
    inh_zh = _inheritance_zh(inheritance)
    mim = phenotype_mim or (v.get("OMIM_id") or "").strip()

    parts = [f"{gene}為"]
    if disease_str:
        parts.append(f"{disease_str}的致病基因之一")
    else:
        parts.append("此疾病的致病基因之一")
    if inh_zh:
        parts.append(f"，其遺傳模式屬於{inh_zh}")
    if mim:
        parts.append(f" (Phenotype MIM number: {mim})")
    return "".join(parts) + "。"


def _picked_clinvar_diseases_for_mito(v: dict, edits: dict) -> list[str]:
    diseases = v.get("clinvar_diseases") or []
    if not isinstance(diseases, list):
        diseases = []
    picked = edits.get("report_diseases_clinvar") or {}
    out: list[str] = []
    if isinstance(picked, dict):
        for idx, disease in enumerate(diseases):
            if picked.get(str(idx)) or picked.get(idx):
                d = str(disease or "").strip()
                if d:
                    out.append(d)
    return out


def _mito_disease_text(v: dict, edits: dict) -> str:
    picked = _picked_clinvar_diseases_for_mito(v, edits)
    if picked:
        return "、".join(picked)
    return ""


def _selected_snv_transcript(v: dict, edits: dict) -> dict:
    options = v.get("transcript_options") or []
    selected = str(edits.get("selected_transcript_key") or "").strip()
    if selected:
        for opt in options:
            if str(opt.get("key") or "") == selected:
                return opt
    default_key = str(v.get("default_transcript_key") or "").strip()
    if default_key:
        for opt in options:
            if str(opt.get("key") or "") == default_key:
                return opt
    return {}


def _snv_tx_field(v: dict, edits: dict, field: str, *fallbacks: str) -> str:
    selected = _selected_snv_transcript(v, edits)
    value = selected.get(field)
    if value not in (None, ""):
        return str(value)
    for key in fallbacks:
        value = v.get(key)
        if value not in (None, ""):
            return str(value)
    value = v.get(field)
    return "" if value in (None, "") else str(value)


def _strip_hgvs_prefix(value: str) -> str:
    text = str(value or "").strip()
    if not text or text in {".", "-", "NA", "N/A"}:
        return ""
    for marker in ("c.", "n.", "m.", "g.", "p."):
        idx = text.find(marker)
        if idx >= 0:
            return text[idx:]
    return text.split(":", 1)[1] if ":" in text else text


def _snv_transcript_label(v: dict, edits: dict) -> str:
    enst = _snv_tx_field(v, edits, "ensembl_transcript")
    refseq = _snv_tx_field(v, edits, "refseq_transcript")
    if enst and refseq and enst != refseq:
        return f"{enst}; {refseq}"
    return refseq or enst or _snv_tx_field(v, edits, "transcript", "MANE_SELECT")


def _snv_variant_block(doc, v: dict, *, tier: str, edits: dict) -> None:
    gene = _snv_tx_field(v, edits, "gene_symbol") or "?"
    tx   = _snv_transcript_label(v, edits)
    _add_paragraph(doc, f"    {gene} ({tx})", bold=True)
    rs    = v.get("rs_id") or v.get("RS_ID") or ""
    struc = _structure_label({
        **v,
        "exon": _snv_tx_field(v, edits, "exon"),
        "intron": _snv_tx_field(v, edits, "intron"),
    })
    hgvs_c = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_C", "hgvs_c"))
    hgvs_p = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_P", "hgvs_p"))
    nuc    = hgvs_c + (f"({hgvs_p})" if hgvs_p else "")
    # 基因型 column stays in English per spec (Heterozygous / Homozygous)
    zyg    = _zygosity_long(v.get("zygosity", ""))
    clnsig = _clinvar_label(v)
    acmg   = _acmg_label(v, edits)

    _ascii_table(doc, columns=[
        ("類別",          5),
        ("基因",          9),
        ("RS ID",         7, "buffered"),
        ("結構",          9),
        ("核苷酸",       14, "hgvs"),
        ("基因型",       13),
        ("ClinVar",      16, "token-buffered"),
        ("ACMG&AMP指引", 12, "token"),
    ], rows=[[tier, gene, rs, struc, nuc, zyg, clnsig, acmg]])

    _add_paragraph(doc, f"    1. {_omim_block_for_snv(v, edits)}")
    _add_paragraph(doc, f"    2. {_patho_sentence(acmg)}")

def _snv_reference_text(v: dict, edits: dict, *, acmg_zh: bool = False) -> str:
    gene = _snv_tx_field(v, edits, "gene_symbol") or "?"
    # Drop the "NM_xxx:" transcript prefix from HGVS values here too —
    # the transcript already appears beside the gene name on the block
    # header, no need to repeat it in the ref text.
    hgvs_c = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_C", "hgvs_c"))
    hgvs_p = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_P", "hgvs_p"))
    nuc    = hgvs_c + (f" ({hgvs_p})" if hgvs_p else "")
    cq_zh  = _consequence_zh(_snv_tx_field(v, edits, "Consequence"))
    af     = _afs_str(v)
    clinvar = _clinvar_label(v)
    acmg    = _acmg_label(v, edits)
    if acmg_zh:
        acmg = _health_pathogenicity_zh(acmg)

    af_sent = (f"該變異位點在族群資料庫 gnomAD 中發生率為 {af}，"
               "顯示其為罕見變異位點") if af \
        else "該變異位點在族群資料庫 gnomAD 中未報導過發生率，顯示其為罕見變異位點"
    clnv_sent = (f"在疾病資料庫 (ClinVar) 中此變異位點被報導為「{clinvar}」"
                 if clinvar and clinvar != "—" else
                 "在疾病資料庫 (ClinVar) 中未被報導過")

    # PubMed sentence intentionally omitted — the new pipeline carries
    # no HGMD reported-cases column.
    return (
        f"    在個案之檢體中，檢測到1個位於基因{gene}的變異位點。"
        f"變異位點 {nuc} 為{cq_zh}"
        f"{af_sent}。{clnv_sent}。"
        "根據美國醫學遺傳學暨基因體學學會 (American College of Medical "
        "Genetics and Genomics) 與分子病理學學會 (Association for "
        f"Molecular Pathology) 於2015年發表之準則，評測此變異位點為「{acmg}」。"
        "此報告僅供參考，臨床判斷仍應以病患的實際狀況為主。"
    )


# ── Mitochondrial ─────────────────────────────────────────────────

def _mito_variant_block(doc, v: dict, *, tier: str, edits: dict) -> None:
    gene  = v.get("gene_symbol") or "?"
    _add_paragraph(doc, f"    {gene} (NC_012920.1)", bold=True)
    hgvs_m = v.get("HGVS_M") or v.get("id") or ""
    aa     = _mito_aa_to_hgvsp(v.get("aa_change") or "")
    nuc    = hgvs_m + (f"({aa})" if aa else "")
    het    = v.get("heteroplasmy")
    try:
        het_s = f"{float(het) * 100:.1f}%" if het not in (None, "", ".") else "—"
    except (TypeError, ValueError):
        het_s = "—"
    clnsig = _clinvar_label(v)
    acmg   = _acmg_label(v, edits)

    # RS ID column dropped — MITOMAP doesn't expose rsIDs in our
    # current adapter. If the pipeline starts emitting one, conditionally
    # add the column back.
    _ascii_table(doc, columns=[
        ("類別",          5),
        ("基因",          9),
        ("核苷酸",       14, "hgvs"),
        ("異質性比例",   13),
        ("ClinVar",      13, "token"),
        ("ACMG&AMP指引", 13, "token"),
    ], rows=[[tier, gene, nuc, het_s, clnsig, acmg]])

    mt_disease = _mito_disease_text(v, edits)
    # 1. 致病基因之一 — ClinVar disease name; 遺傳模式 fixed to 粒線體遺傳; no MIM
    if mt_disease:
        _add_paragraph(doc, f"    1. {gene}為{mt_disease}的致病基因之一，其遺傳模式屬於粒線體遺傳。")
    else:
        _add_paragraph(doc, f"    1. {gene}為粒線體疾病之相關基因，其遺傳模式屬於粒線體遺傳。")
    _add_paragraph(doc, f"    2. {_patho_sentence(acmg)}")
    _add_paragraph(doc, "    3. 此檢驗結果須由臨床醫師來分析粒線體DNA基因檢驗結果與受檢者臨床症狀的相關性"
                        "並考量家族史的關聯性。此外，粒線體DNA基因變異有組織的特異性，粒線體DNA基因致病變異的"
                        "異質性（heteroplasmy）在不同組織間會有差異。")
    _add_paragraph(doc, "    4. 對於所有粒線體DNA基因變異所造成的臨床症狀表現的差異，可能取決於下列三項因素："
                        "（1）異質性，即致病變異及正常粒線體 DNA 存在量的比例（2）變異的粒線體 DNA 於組織的"
                        "分布情形（3）限界效應（Threshold effect），即每種組織對於氧化壓力代謝影響的易受性。"
                        "由於患者情況各異，此檢驗結果須由臨床醫師判讀檢驗結果與受檢者臨床症狀的相關性。")
def _mito_reference_text(v: dict, edits: dict) -> str:
    gene = v.get("gene_symbol") or "?"
    hgvs_m = v.get("HGVS_M") or v.get("id") or ""
    aa = _mito_aa_to_hgvsp(v.get("aa_change") or "")
    nuc = hgvs_m + (f"({aa})" if aa else "")
    cq_zh = _consequence_zh(v.get("consequence", ""))
    clinvar = _clinvar_label(v)
    acmg    = _acmg_label(v, edits)
    clnv_sent = (f"在疾病資料庫 (ClinVar) 中此變異位點被報導為「{clinvar}」"
                 if clinvar and clinvar != "—"
                 else "在疾病資料庫 (ClinVar) 中未被報導過")
    return (
        f"    在個案之檢體中，檢測到1個位於基因{gene}的變異位點。"
        f"變異位點 {nuc} 為{cq_zh}"
        f"{clnv_sent}。"
        "根據美國醫學遺傳學暨基因體學學會 (American College of Medical "
        "Genetics and Genomics) 與分子病理學學會 (Association for "
        "Molecular Pathology) 於2015年發表之準則，並參照ClinGen Mitochondrial "
        "Disease Expert Panel於 2020 年發布之粒線體DNA變異判讀專用ACMG/AMP標準"
        f"進行評估，評測此變異位點為「{acmg}」。"
        "此報告僅供參考，臨床判斷仍應以病患的實際狀況為主。"
    )


# ── CNV / SV ──────────────────────────────────────────────────────

def _build_coords(v: dict) -> str:
    """Prefer the adapter's `coords` field; fall back to constructing
    `chrN:POS-END` from CHROM / POS / END when it's missing or blank.
    Always include the `chr` prefix on the chromosome name (UCSC
    convention; matches the UI header `[GRCh38] chr1:...`).
    """
    coords = (v.get("coords") or "").strip()
    if not coords:
        chrom = (v.get("CHROM") or "").strip()
        pos   = v.get("POS")
        end   = v.get("END")
        if chrom and pos is not None and end is not None:
            coords = f"{chrom}:{pos}-{end}"
    if coords and not coords.startswith("chr"):
        coords = "chr" + coords
    return coords


def _location_zh(g: dict) -> str:
    """AnnotSV's per-gene `Location` field (e.g. 'txStart-txEnd',
    'txStart-intron3', 'exon2-exon8') → 中文 phrase for the report.

    Whole-gene coverage ('txStart-txEnd') always returns 「整個區域」.
    Partial coverage parses each endpoint:
      txStart → 基因起始, txEnd → 基因末端,
      exonN   → "Exon N 區域", intronN → "Intron N 區域"
    Joined with 「至」. A leading ASCII letter on the second piece gets
    a leading space ("基因起始至 Intron 3 區域") so Latin tokens don't
    sit right against the Chinese 至 separator.
    Unrecognised endpoints pass through verbatim so weird AnnotSV
    values stay visible for review.
    """
    loc = (g.get("location") or "").strip() if isinstance(g, dict) else ""
    if not loc or loc == "txStart-txEnd":
        return "整個區域"
    parts = loc.split("-", 1)
    if len(parts) != 2:
        return loc

    def piece(p: str) -> str:
        p = p.strip()
        if p == "txStart": return "基因起始"
        if p == "txEnd":   return "基因末端"
        m = re.match(r"^exon(\d+)$", p, re.IGNORECASE)
        if m: return f"Exon {m.group(1)} 區域"
        m = re.match(r"^intron(\d+)$", p, re.IGNORECASE)
        if m: return f"Intron {m.group(1)} 區域"
        return p

    p1 = piece(parts[0])
    p2 = piece(parts[1])
    sep = "至" + (" " if p2 and "A" <= p2[0] <= "Z" else "")
    return f"{p1}{sep}{p2}"


def _gene_loc_phrase(gname: str, loc_zh: str) -> str:
    """Build「GENE 基因之 …」 / 「GENE 基因之 Exon X 區域」 — always
    with 之, with a space inserted when the location phrase starts
    with a Latin letter so "基因之 Exon" reads cleanly."""
    sep = "之" + (" " if loc_zh and "A" <= loc_zh[0] <= "Z" else "")
    return f"{gname} 基因{sep}{loc_zh}"


def _omim_genes(v: dict) -> list[dict]:
    """Filter the variant's `genes` list to those carrying an OMIM_ID —
    the only ones worth surfacing in the diagnostic report."""
    out = []
    for g in (v.get("genes") or []):
        if not isinstance(g, dict):
            continue
        if (g.get("omim_id") or "").strip():
            out.append(g)
    return out


def _coords_str(v: dict, is_wgs: bool) -> str:
    """WES uses imprecise breakpoint notation `(?_start)_(end_?)`;
    WGS uses the precise start_end form. The adapter's `coords` is the
    raw AnnotSV CHROM:POS-END so we annotate accordingly.
    """
    coords = _build_coords(v)
    if is_wgs or not coords:
        return coords
    # WES form
    m = re.match(r"(chr[\dXYM]+):(\d+)-(\d+)", coords)
    if not m:
        return coords
    ch, s, e = m.group(1), m.group(2), m.group(3)
    return f"{ch}:(?_{s})_({e}_?)"


def _sv_kind_zh(v: dict) -> str:
    return _CNV_KIND_ZH.get((v.get("sv_type") or "").upper(), "變異")


def _cnv_report_disease(edits: dict) -> str:
    """Reviewer-entered CNV/SV disease label for the formal report."""
    return str(edits.get("disease") or "").strip()


def _cnv_variant_block(doc, v: dict, *, tier: str, is_wgs: bool,
                       edits: dict) -> None:
    coords    = _coords_str(v, is_wgs)
    sv_upper  = (v.get("sv_type") or "").upper()
    cn        = v.get("copy_number")
    cn_s      = "—" if cn in (None, "", ".") else str(cn)
    # 基因型 only meaningful for DEL — DUP / INS / INV reuse the column
    # for "—" so the reviewer doesn't read het/hom on a copy gain.
    raw_zyg   = (v.get("zygosity", "") or "").strip()
    if sv_upper == "DEL":
        zyg = _zygosity_long(raw_zyg) or "—"
    else:
        zyg = "—"
    acmg      = _acmg_label(v, edits)
    chrom     = (v.get("CHROM") or "").replace("chr", "")
    coords_disp = coords.split(":", 1)[1] if ":" in coords else coords

    sv_tag = "del" if sv_upper == "DEL" else ("dup" if sv_upper == "DUP" else "")
    _add_paragraph(doc, f"    [GRCh38] {coords}{sv_tag}", bold=True)
    _ascii_table(doc, columns=[
        ("類別",          5),
        ("染色體",        7),
        ("變異位置",     26, "buffered"),
        ("拷貝數",        7),
        ("基因型",       13),
        ("ACMG&AMP指引", 13, "token"),
    ], rows=[[tier, chrom, coords_disp, cn_s, zyg, acmg]])

    # 1. 片段位置描述 — only surface OMIM-tagged genes (per spec);
    # decide single-gene vs multi-gene template by how many are left.
    chrom_num  = (v.get("CHROM") or "").replace("chr", "")
    omim_genes = _omim_genes(v)

    # Reviewer can prune which genes appear via edits.report_genes
    # (frontend toggle on the CNV/SV card). Honour the prune only when
    # at least one OMIM gene survives.
    report_genes = edits.get("report_genes") or {}
    if isinstance(report_genes, dict):
        kept = [g for g in omim_genes if report_genes.get(g.get("gene"))]
        if kept:
            omim_genes = kept

    kind_zh = _CNV_KIND_ZH.get((v.get("sv_type") or "").upper(), "變異")

    next_idx = 2
    disease_override = _cnv_report_disease(edits)
    if len(omim_genes) == 1:
        g = omim_genes[0]
        gname = g.get("gene") or "?"
        loc_zh = _location_zh(g)
        _add_paragraph(doc, f"    1. 此片段位於第 {chrom_num} 號染色體上 {_gene_loc_phrase(gname, loc_zh)}。")
        # 2. OMIM phenotype + inheritance, per-gene
        ph, ph_inheritance, phenotype_mim = _disease_info(
            disease_override or (g.get("omim_phenotype") or "").strip()
        )
        inh = _inheritance_zh(
            ph_inheritance or g.get("omim_inheritance", "") or ""
        )
        mim = phenotype_mim or (g.get("omim_id") or "").strip()
        bits = [f"{gname}為"]
        bits.append(f"{ph}的致病基因之一" if ph else "此疾病的致病基因之一")
        if inh: bits.append(f"，其遺傳模式屬於{inh}")
        if mim: bits.append(f" (Phenotype MIM number: {mim})")
        _add_paragraph(doc, f"    2. {''.join(bits)}。")
        next_idx = 3
    elif len(omim_genes) > 1:
        names = [g.get("gene", "") for g in omim_genes[:10] if g.get("gene")]
        disease_suffix = f"，與 {disease_override} 相關" if disease_override else ""
        _add_paragraph(doc, f"    1. 此片段位於第 {chrom_num} 號染色體上，"
                            f"包含 {', '.join(names)} 等 OMIM 疾病基因"
                            f"{disease_suffix}。")
    else:
        disease_suffix = f"，與 {disease_override} 相關" if disease_override else ""
        _add_paragraph(doc, f"    1. 此片段位於第 {chrom_num} 號染色體上，"
                            "未涵蓋 OMIM 疾病相關基因"
                            f"{disease_suffix}。")

    _add_paragraph(doc, f"    {next_idx}. {_patho_sentence(acmg)}")
    next_idx += 1
    if not is_wgs:
        _add_paragraph(doc, f"    {next_idx}. 由於此檢驗技術為全外顯子定序，"
                            "若缺失片段之斷點(Breakpoints)發生於內含子(Intron) ，"
                            "則無法明確判別起始及末端位置。")

def _cnv_reference_text(v: dict, edits: dict, omim_genes: list[dict],
                        kind_zh: str, is_wgs: bool) -> str:
    coords = _coords_str(v, is_wgs) or _build_coords(v)
    # Zygosity in the sentence only when DEL — DUP / INS / INV don't
    # carry meaningful zygosity (a triploid duplication isn't "het").
    sv_upper = (v.get("sv_type") or "").upper()
    if sv_upper == "DEL":
        zyg_text = _zygosity_zh(v.get("zygosity", "")) or "異合子"
        zyg_phrase = f"之{zyg_text}"
    else:
        zyg_phrase = "之"
    acmg    = _acmg_label(v, edits)

    if len(omim_genes) == 1:
        g = omim_genes[0]
        gname = g.get("gene") or "?"
        loc_zh = _location_zh(g)
        span_desc = f"此段{kind_zh}涵蓋 {_gene_loc_phrase(gname, loc_zh)}"
    elif len(omim_genes) > 1:
        names = [g.get("gene", "") for g in omim_genes[:10] if g.get("gene")]
        span_desc = f"此段{kind_zh}涵蓋 {', '.join(names)} 等 OMIM 疾病基因"
    else:
        span_desc = f"此段{kind_zh}未涵蓋 OMIM 疾病相關基因"

    disease = _cnv_report_disease(edits)
    disease_sent = f"此變異與「{disease}」相關。" if disease else ""
    return (
        f"    在個案之檢體中，檢測到位於 {coords} {zyg_phrase}片段{kind_zh}變異，"
        f"{span_desc}。"
        f"{disease_sent}"
        "根據美國醫學遺傳學暨基因體學學會 (American College of Medical "
        "Genetics and Genomics) 與分子病理學學會 (Association for Molecular "
        "Pathology) 於2015年發表之準則，並參照ClinGen 及Riggs等人於 2020 年發布之"
        f"拷貝數變異判讀專用ACMG/ClinGen技術標準進行評估，評測此變異位點為「{acmg}」。"
        "此報告僅供參考，臨床判斷仍應以病患的實際狀況為主。"
    )


# ── §四 方法、§五 注釋 ───────────────────────────────────────────

def _section_methods(doc, test_type: str, *, health: bool = False) -> None:
    is_wgs = (test_type or "").upper() == "WGS"
    seq    = "Illumina NovaSeq X Plus" if is_wgs else "Illumina NextSeq 2000"
    depth  = "27X" if is_wgs and health else ("30X" if is_wgs else "50X")
    _add_paragraph(doc, "四、檢測方法說明")
    _add_paragraph(doc, f"  1. 本次檢測使用次世代定序儀分析 ({seq})。")
    _add_paragraph(doc, "  2. 本次檢測變異位點的錯誤率 ≦ 0.1% (Phred-scaled Q score ≧ 30)。")
    _add_paragraph(doc, f"  3. 本次檢測平均定序深度 ≧ {depth}。")
    if health:
        _add_paragraph(doc, "  4. 本檢測僅能檢測出基因內單一核苷酸變異 (single nucleotide variant) 、"
                            "小片段的缺失或插入 (small indel)，無法檢測出拷貝數變異 "
                            "(copy number variant)、轉位 (translocation)、"
                            "倒轉 (inversion) 或其他複雜性結構變異 "
                            "(complex structural variant)、組織特異性的鑲嵌 (tissue-specific mosaicism) "
                            "以及未包含在本次定序範圍之區域。")
        _add_paragraph(
            doc,
            "  5. 藥物基因體學分析中，CYP2D6 基因型判定會納入該基因之拷貝數變異 "
            "(copy number variation) 分析結果；此項專一性分析僅用於 CYP2D6 藥物基因體學判讀，"
            "不代表本檢測已涵蓋其他基因之拷貝數變異。",
        )
        next_note_number = 6
    else:
        _add_paragraph(doc, "  4. 本檢測僅能檢測出基因內單一核苷酸變異 (single nucleotide variant) 、"
                            "小片段的缺失或插入 (small indel)及部分拷貝數變異 (copy number variant)，"
                            "無法檢測出轉位 (translocation)、倒轉 (inversion) 或其他複雜性結構變異 "
                            "(complex structural variant)、組織特異性的鑲嵌 (tissue-specific mosaicism) "
                            "以及未包含在本次定序範圍之區域。")
        next_note_number = 5
    if is_wgs:
        _add_paragraph(doc, f"  {next_note_number}. 本實驗方法以次世代方法定序粒線體DNA基因序列，"
                            "變異點位判讀之cut-off值定為5%異質性（heteroplasmy）。")
        next_note_number += 1
    _add_paragraph(doc, f"  {next_note_number}. 本檢測報告僅供醫療專業人員參考，需配合其他相關臨床資料與家族成員之相關檢驗。"
                        "依衛福部規定，目前次世代定序分子遺傳診斷皆屬研究性質。")
    _blank(doc)


def _section_annotations(doc, sample: dict, gene_list_mode: str) -> None:
    _add_paragraph(doc, "五、檢測結果注釋")
    _add_paragraph(doc, "  1. 本檢測結果比對參考序列為人類hg38版本。")
    _add_paragraph(doc, f"  2. ClinVar及ACMG&AMP指引：引用ClinVar資料庫截至{CLINVAR_DATE_HUMN}更新的註解，"
                        "及美國醫學遺傳學暨基因體學學會 (ACMG) 與分子病理學學會 (AMP) 2015年頒佈的指引，"
                        "並且主要列入致病(Pathogenic) 及疑似致病 (Likely pathogenic) 變異；"
                        "其他類別變異經醫師判斷認為與疾病相關時亦可列入。")
    _add_paragraph(doc, "  3. 參考資料:")
    _add_paragraph(doc, f"     a. 疾病資料庫: OMIM、ClinVar ({CLINVAR_DATE})")
    _add_paragraph(doc, "     b. 族群資料庫: gnomAD (v4.1 genome)")
    _add_paragraph(doc, "     c. 序列資料庫: RefSeqGene (105.20220307)")
    _add_paragraph(doc, "  4. 本次檢測基因包括")
    _render_gene_list(doc, sample, gene_list_mode)


def _hpo_label_for(hid: str) -> str:
    term = hpo_ontology.get(hid)
    return term.name if term else hid


def _genes_for_term_or_panel(key: str) -> list[str]:
    """phenotype_scorer._HPO_TO_GENES is the union of HPO→gene map and
    every panel's gene set (HPO ids + panel names live in the same dict)."""
    genes = []
    for raw in phenotype_scorer._HPO_TO_GENES.get(key, set()):
        gene, _hid = panel_deadzone.canonical_gene_symbol(raw)
        if panel_deadzone.is_disease_associated_gene(gene):
            genes.append(gene)
    return sorted(set(genes))


def _gene_list_label(gene: str, test_type: str) -> str:
    return gene


def _add_dead_zone_gene_list_note(doc, threshold: int) -> None:
    _blank(doc)
    _add_paragraph(doc, f"註：括號中標示之 exon 為 cohort dead-zone，代表該 exon coverage 低於本檢測判讀門檻（<{threshold}X）。")


def _render_gene_list(doc, sample: dict, mode: str) -> None:
    """mode = 'grouped' → one paragraph per HPO term / panel,
       mode = 'merged'  → single deduped list.
    """
    hpo_rows: list = sample.get("patient_phenotype") or []
    # selected_panels is a list of {name, weight} dicts (see
    # phenotype_io.parse) — pull the name field.
    panel_entries: list = sample.get("selected_panels") or []
    test_type = ((sample.get("meta") or {}).get("Test") or "WES").upper()
    include_docx_dead_zone = False

    # Build [(display_name, [genes...])] preserving order.
    sections: list[tuple[str, list[str]]] = []
    for r in hpo_rows:
        hid = r.get("phenotype") or ""
        label = r.get("label") or _hpo_label_for(hid)
        if not hid: continue
        genes = _genes_for_term_or_panel(hid)
        sections.append((label, genes))
    for entry in panel_entries:
        pname = entry.get("name") if isinstance(entry, dict) else str(entry)
        if not pname: continue
        genes = _genes_for_term_or_panel(pname)
        sections.append((pname, genes))

    if not sections:
        _add_paragraph(doc, "    （未設定 HPO / panel — 無檢測基因清單）")
        return

    if mode == "merged":
        merged: set[str] = set()
        for _, gs in sections:
            merged |= set(gs)
        gene_str = ", ".join(_gene_list_label(g, test_type) for g in sorted(merged))
        _add_paragraph(doc, gene_str)
        if include_docx_dead_zone:
            _add_dead_zone_gene_list_note(doc, threshold)
        return

    # grouped (default)
    for idx, (name, gs) in enumerate(sections):
        if idx:
            _blank(doc)
        _add_paragraph(doc, f"{name}:")
        if gs:
            _add_paragraph(doc, ", ".join(_gene_list_label(g, test_type) for g in gs))
        else:
            _add_paragraph(doc, "（無對應基因）")
    if include_docx_dead_zone:
        _add_dead_zone_gene_list_note(doc, threshold)


# ── Health screening DOCX ─────────────────────────────────────────

_NO_HEALTH_VARIANT_TEXT = "於本次檢測之基因中，未檢出疾病資料庫中已收錄之致病性或疑似致病性變異。"
_NO_HEALTH_CARRIER_VARIANT_TEXT = (
    "於本次檢測之基因中，未檢出疾病資料庫中已收錄且符合帶因者狀態之致病性或疑似致病性變異。"
)

_HEALTH_NEGATIVE_LIMITATION = (
    "於本次檢測之基因中未檢出疾病資料庫中已收錄之致病性或疑似致病性變異，並不能排除受檢者仍具有"
    "相關遺傳性疾病風險。"
    "未被偵測或未被回報之情形可能包括：資料庫尚未收錄之新變異、目前證據不足或分類不一致之變異、"
    "拷貝數變異、結構變異、重複序列變異、低比例鑲嵌、低覆蓋或比對困難區域，以及本檢測方法未涵蓋"
    "之變異型態。變異分類與臨床意義可能隨資料庫與醫學知識更新而改變。若受檢者已有相關症狀或明確"
    "家族史，仍建議接受專科評估。"
)

_HEALTH_ACMG_CAUTION = (
    f"本檢測依據美國醫學遺傳學暨基因體學學會（ACMG）次發現基因清單第 3.3 版（2025 年發表），"
    f"分析與 ACMG SF 3.3 所列遺傳性疾病相關的風險基因，並涵蓋疾病資料庫 ClinVar（版本 "
    f"{CLINVAR_DATE}）中已收錄之致病性或疑似致病性變異。檢測結果代表於本次檢測範圍及技術限制內"
    "所辨識的變異，不代表受檢者目前已罹患相關疾病，也不能完全排除未來發病或其他遺傳性疾病"
    "的可能性。疾病風險仍須綜合個人病史、家族史、臨床檢查及其他檢驗結果評估。"
)

_HEALTH_PGX_CAUTION = (
    "本報告之藥物基因體學結果係依本檢測涵蓋之基因與變異位點，推估受檢者對部分藥物之代謝能力、"
    "療效或不良反應風險，僅供合格醫療專業人員作為處方與用藥評估參考。藥物反應亦受年齡、性別、"
    "體重、肝腎功能、共病、懷孕狀態、飲食、吸菸、併用藥物、藥物交互作用及實際用藥適應症影響。"
    "受檢者不應自行停藥、換藥或調整劑量；任何用藥變更應由處方醫師、臨床藥師或臨床藥物基因體學"
    "門診評估。本報告可能包含受檢者目前未使用的藥物。相關結果可供未來處方時參考，不代表目前"
    "需要使用、停用或更換任何藥物。"
)

_HEALTH_PGX_RESOURCE_INTRO = (
    "以上簡要判讀說明未提供完整的劑量調整指示。可供臨床處方決策參考的藥物基因體資訊，"
    "可於如下之 CPIC 最新臨床指引、藥物之 FDA 核准仿單及目前的藥物基因體學相關生物標記表中查詢。"
)

_HEALTH_PGX_RESOURCE_CLOSING = (
    "調整藥物劑量前，或需要進一步資訊時，請務必諮詢臨床醫師或臨床藥理專業人員。"
)

_HEALTH_PGX_RESOURCES = (
    (
        "CPIC 最新臨床指引",
        "https://cpicpgx.org/guidelines/",
    ),
    (
        "FDA 核准藥品仿單",
        "https://www.accessdata.fda.gov/scripts/cder/daf/index.cfm",
    ),
    (
        "FDA 藥物基因體關聯表",
        "https://www.fda.gov/medical-devices/precision-medicine/table-pharmacogenetic-associations",
    ),
    (
        "FDA 藥品仿單之藥物基因體生物標記表",
        "https://www.fda.gov/drugs/science-and-research-drugs/table-pharmacogenomic-biomarkers-drug-labeling",
    ),
)

_HEALTH_SECTION_ORDER = [
    ("acmg_sf", "第一類：與疾病風險相關之致病性或疑似致病性變異位點"),
    ("lipid_fh", "第二類：血脂相關基因"),
    ("hereditary_cancer", "第三類：腫瘤相關基因"),
    ("stroke", "第四類：中風相關基因"),
    ("carrier", "第五類：帶因者篩查"),
    ("proactive", "第六類：主動篩查"),
    ("pgx", "藥物基因體學"),
]

_HEALTH_ACMG_GENE_LIST_TITLE = (
    "第一類：ACMG疾病風險基因（參考美國醫學遺傳學暨基因體學學會 "
    "(American College of Medical Genetics and Genomics) 於 2025 年所公告之"
    "次發現 (Secondary findings) 基因清單 v3.3 版；PMID: 40568962）"
)

_HEALTH_PGX_GENE_LIST_TITLE = (
    "藥物基因體學（參考 Clinical Pharmacogenetics Implementation Consortium (CPIC) "
    "Level A gene–drug pairs）"
)

_PGX_CPIC_LEVEL_A_GENES = (
    "ABCG2", "CACNA1S", "CFTR", "CYP2B6", "CYP2C19", "CYP2C9",
    "CYP2D6", "CYP3A5", "CYP4F2", "DPYD", "G6PD", "HLA-A", "HLA-B",
    "MT-RNR1", "NAT2", "NUDT15", "RYR1", "SLCO1B1", "TPMT",
    "UGT1A1", "VKORC1",
)
_PGX_CPIC_LEVEL_A_SET = set(_PGX_CPIC_LEVEL_A_GENES)
_PGX_FULL_RECOMMENDATION_COLUMNS = [
    ("藥物", 23, "word-buffered"),
    ("基因與表型", 18, "word-buffered"),
    ("CPIC/FDA 建議", 43, "word-buffered"),
]


def _meta_value(meta: dict, key: str) -> str:
    value = meta.get(key)
    return "" if value in (None, "") else str(value)


def _display_lis_id(value: str) -> str:
    text = str(value or "").strip()
    for suffix in ("-dragen", "-nckuh", "-inhouse"):
        if text.lower().endswith(suffix):
            return text[: -len(suffix)]
    return text


def _health_meta_row(row: list[tuple[str, str]]) -> str:
    widths = (32, 25, 25)
    cells = []
    for (label, value), width in zip(row, widths):
        cells.append(_pad_right(f"{label}: {value or '—'}", width))
    return " ".join(cells).rstrip()


def _section_health_patient_header(doc, meta: dict) -> None:
    """Basic case metadata above the health-report sections."""
    _add_paragraph(doc, "國立成功大學醫學院附設醫院", bold=True, align="center")
    _add_paragraph(doc, "<<基因醫學部基因檢測分析研究報告>>", bold=True, align="center")
    _blank(doc)
    rows = [
        (
            ("檢體編號", _display_lis_id(_meta_value(meta, "LIS_ID"))),
            ("病歷號", _meta_value(meta, "MRN")),
            ("簽收日期", _meta_value(meta, "SignReceivedAt")),
        ),
        (
            ("姓名", _meta_value(meta, "Name")),
            ("性別", _meta_value(meta, "Sex")),
            ("出生日期", _meta_value(meta, "DOB")),
        ),
    ]
    _add_paragraph(doc, "受檢者資料", bold=True)
    for row in rows:
        _add_paragraph(doc, _health_meta_row(list(row)))
    _blank(doc)

_ACMG_SF_GROUPS = [
    {
        "key": "lipid",
        "title": "血脂相關基因",
        "genes": ["APOB", "LDLR", "PCSK9"],
    },
    {
        "key": "tumor",
        "title": "腫瘤相關基因",
        "genes": [
            "APC", "BMPR1A", "BRCA1", "BRCA2", "MAX", "MEN1", "MLH1",
            "MSH2", "MSH6", "MUTYH", "NF2", "PALB2", "PMS2", "PTEN",
            "RB1", "RET", "SDHAF2", "SDHB", "SDHC", "SDHD", "SMAD4",
            "STK11", "TMEM127", "TP53", "TSC1", "TSC2", "VHL", "WT1",
        ],
    },
    {
        "key": "cardiovascular",
        "title": (
            "心血管疾病相關基因，包含心肌病變相關基因（ACTC1, BAG3, DES, DSC2, "
            "DSG2, DSP, FLNC, LMNA, MYBPC3, MYH7, MYL2, MYL3, PKP2, PLN, "
            "PRKAG2, RBM20, TMEM43, TNNC1, TNNI3, TNNT2, TPM1, TTN）、"
            "心律不整相關基因（CALM1, CALM2, CALM3, CASQ2, KCNH2, KCNQ1, "
            "RYR2, SCN5A, TRDN）、主動脈及血管疾病相關基因（ACTA2, COL3A1, "
            "FBN1, MYH11, SMAD3, TGFBR1, TGFBR2, ACVRL1, ENG）"
        ),
        "genes": [
            "ACTC1", "BAG3", "DES", "DSC2", "DSG2", "DSP", "FLNC",
            "LMNA", "MYBPC3", "MYH7", "MYL2", "MYL3", "PKP2", "PLN",
            "PRKAG2", "RBM20", "TMEM43", "TNNC1", "TNNI3", "TNNT2",
            "TPM1", "TTN", "CALM1", "CALM2", "CALM3", "CASQ2", "KCNH2",
            "KCNQ1", "RYR2", "SCN5A", "TRDN", "ACTA2", "COL3A1",
            "FBN1", "MYH11", "SMAD3", "TGFBR1", "TGFBR2", "ACVRL1",
            "ENG",
        ],
    },
    {
        "key": "metabolic_endocrine",
        "title": "代謝與內分泌疾病相關基因",
        "genes": ["ABCD1", "ATP7B", "BTD", "CYP27A1", "GAA", "GLA", "HFE", "HNF1A", "OTC"],
    },
    {
        "key": "anesthesia",
        "title": "麻醉用藥風險相關基因",
        "genes": ["CACNA1S", "RYR1"],
    },
    {
        "key": "other",
        "title": "其它基因",
        "genes": ["RPE65", "TTR"],
    },
]


def _health_acmg_group_label(group: dict, index: int) -> str:
    """Return the legacy ACMG subgroup label now shown in the gene appendix."""
    if group.get("key") == "cardiovascular":
        return f"{index}. {group['title']}"
    genes = ", ".join(group.get("genes") or [])
    return f"{index}. {group['title']}，包含 {genes}"


_ACMG_SF_DISEASES = {
    "ABCD1": [{"disease": "X-linked adrenoleukodystrophy", "inheritance": "XL"}],
    "ACTA2": [{"disease": "Familial thoracic aortic aneurysm", "inheritance": "AD"}],
    "ACTC1": [{"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}],
    "ACVRL1": [{"disease": "Hereditary hemorrhagic telangiectasia", "inheritance": "AD"}],
    "APC": [{"disease": "Familial adenomatous polyposis", "inheritance": "AD"}],
    "APOB": [{"disease": "Familial hypercholesterolemia", "inheritance": "AD"}],
    "ATP7B": [{"disease": "Wilson disease", "inheritance": "AR"}],
    "BAG3": [{"disease": "Dilated cardiomyopathy", "inheritance": "AD"}, {"disease": "Myofibrillar myopathy", "inheritance": "AD"}],
    "BMPR1A": [{"disease": "Juvenile polyposis syndrome", "inheritance": "AD"}],
    "BRCA1": [{"disease": "Hereditary breast and ovarian cancer", "inheritance": "AD"}],
    "BRCA2": [{"disease": "Hereditary breast and ovarian cancer", "inheritance": "AD"}],
    "BTD": [{"disease": "Biotinidase deficiency", "inheritance": "AR"}],
    "CACNA1S": [{"disease": "Malignant hyperthermia", "inheritance": "AD"}],
    "CALM1": [{"disease": "Long-QT syndrome type 14", "inheritance": "AD"}, {"disease": "Catecholaminergic polymorphic ventricular tachycardia", "inheritance": "AD"}],
    "CALM2": [{"disease": "Long-QT syndrome type 15", "inheritance": "AD"}, {"disease": "Catecholaminergic polymorphic ventricular tachycardia", "inheritance": "AD"}],
    "CALM3": [{"disease": "Long-QT syndrome type 16", "inheritance": "AD"}, {"disease": "Catecholaminergic polymorphic ventricular tachycardia", "inheritance": "AD"}],
    "CASQ2": [{"disease": "Catecholaminergic polymorphic ventricular tachycardia", "inheritance": "AR"}],
    "COL3A1": [{"disease": "Ehlers-Danlos syndrome, vascular type", "inheritance": "AD"}],
    "CYP27A1": [{"disease": "Cerebrotendinous xanthomatosis", "inheritance": "AR"}],
    "DES": [{"disease": "Dilated cardiomyopathy", "inheritance": "AD"}, {"disease": "Myofibrillar myopathy", "inheritance": "AD"}],
    "DSC2": [{"disease": "Arrhythmogenic right ventricular cardiomyopathy", "inheritance": "AD"}],
    "DSG2": [{"disease": "Arrhythmogenic right ventricular cardiomyopathy", "inheritance": "AD"}],
    "DSP": [{"disease": "Arrhythmogenic right ventricular cardiomyopathy", "inheritance": "AD"}, {"disease": "Dilated cardiomyopathy", "inheritance": "AD"}],
    "ENG": [{"disease": "Hereditary hemorrhagic telangiectasia", "inheritance": "AD"}],
    "FBN1": [{"disease": "Marfan syndrome", "inheritance": "AD"}],
    "FLNC": [{"disease": "Dilated cardiomyopathy", "inheritance": "AD"}, {"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}, {"disease": "Myofibrillar myopathy", "inheritance": "AD"}],
    "GAA": [{"disease": "Pompe disease", "inheritance": "AR"}],
    "GLA": [{"disease": "Fabry disease", "inheritance": "XL"}],
    "HFE": [{"disease": "Hereditary hemochromatosis (c.845G>A; p.C282Y homozygotes only)", "inheritance": "AR"}],
    "HNF1A": [{"disease": "Maturity-Onset of Diabetes of the Young", "inheritance": "AD"}],
    "KCNH2": [{"disease": "Long-QT syndrome type 2", "inheritance": "AD"}],
    "KCNQ1": [{"disease": "Long-QT syndrome type 1", "inheritance": "AD"}],
    "LDLR": [{"disease": "Familial hypercholesterolemia", "inheritance": "AD"}],
    "LMNA": [{"disease": "Dilated cardiomyopathy", "inheritance": "AD"}],
    "MAX": [{"disease": "Hereditary paraganglioma-pheochromocytoma syndrome", "inheritance": "AD"}],
    "MEN1": [{"disease": "Multiple endocrine neoplasia type 1", "inheritance": "AD"}],
    "MLH1": [{"disease": "Lynch syndrome", "inheritance": "AD"}],
    "MSH2": [{"disease": "Lynch syndrome", "inheritance": "AD"}],
    "MSH6": [{"disease": "Lynch syndrome", "inheritance": "AD"}],
    "MUTYH": [{"disease": "MUTYH-associated polyposis", "inheritance": "AR"}],
    "MYBPC3": [{"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}],
    "MYH11": [{"disease": "Familial thoracic aortic aneurysm", "inheritance": "AD"}],
    "MYH7": [{"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}, {"disease": "Dilated cardiomyopathy", "inheritance": "AD"}],
    "MYL2": [{"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}],
    "MYL3": [{"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}],
    "NF2": [{"disease": "NF2-related schwannomatosis", "inheritance": "AD"}],
    "OTC": [{"disease": "Ornithine transcarbamylase deficiency", "inheritance": "XL"}],
    "PALB2": [{"disease": "Hereditary breast cancer", "inheritance": "AD"}],
    "PCSK9": [{"disease": "Familial hypercholesterolemia", "inheritance": "AD"}],
    "PKP2": [{"disease": "Arrhythmogenic right ventricular cardiomyopathy", "inheritance": "AD"}],
    "PLN": [{"disease": "Dilated cardiomyopathy", "inheritance": "AD"}],
    "PMS2": [{"disease": "Lynch syndrome", "inheritance": "AD"}],
    "PRKAG2": [{"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}],
    "PTEN": [{"disease": "PTEN hamartoma tumor syndrome", "inheritance": "AD"}],
    "RB1": [{"disease": "Retinoblastoma", "inheritance": "AD"}],
    "RBM20": [{"disease": "Dilated cardiomyopathy", "inheritance": "AD"}],
    "RET": [{"disease": "Familial medullary thyroid cancer", "inheritance": "AD"}, {"disease": "Multiple endocrine neoplasia type 2A", "inheritance": "AD"}, {"disease": "Multiple endocrine neoplasia type 2B", "inheritance": "AD"}],
    "RPE65": [{"disease": "RPE65-related retinopathy", "inheritance": "AR"}],
    "RYR1": [{"disease": "Malignant hyperthermia", "inheritance": "AD"}],
    "RYR2": [{"disease": "Catecholaminergic polymorphic ventricular tachycardia", "inheritance": "AD"}],
    "SCN5A": [{"disease": "Long QT syndrome type 3", "inheritance": "AD"}, {"disease": "Brugada syndrome", "inheritance": "AD"}, {"disease": "Dilated cardiomyopathy", "inheritance": "AD"}],
    "SDHAF2": [{"disease": "Hereditary paraganglioma-pheochromocytoma syndrome", "inheritance": "AD"}],
    "SDHB": [{"disease": "Hereditary paraganglioma-pheochromocytoma syndrome", "inheritance": "AD"}],
    "SDHC": [{"disease": "Hereditary paraganglioma-pheochromocytoma syndrome", "inheritance": "AD"}],
    "SDHD": [{"disease": "Hereditary paraganglioma-pheochromocytoma syndrome", "inheritance": "AD"}],
    "SMAD3": [{"disease": "Loeys-Dietz syndrome", "inheritance": "AD"}],
    "SMAD4": [{"disease": "Juvenile polyposis syndrome", "inheritance": "AD"}, {"disease": "Hereditary hemorrhagic telangiectasia", "inheritance": "AD"}],
    "STK11": [{"disease": "Peutz-Jeghers syndrome", "inheritance": "AD"}],
    "TGFBR1": [{"disease": "Loeys-Dietz syndrome", "inheritance": "AD"}],
    "TGFBR2": [{"disease": "Loeys-Dietz syndrome", "inheritance": "AD"}],
    "TMEM127": [{"disease": "Hereditary paraganglioma-pheochromocytoma syndrome", "inheritance": "AD"}],
    "TMEM43": [{"disease": "Arrhythmogenic right ventricular cardiomyopathy", "inheritance": "AD"}],
    "TNNC1": [{"disease": "Dilated cardiomyopathy", "inheritance": "AD"}],
    "TNNI3": [{"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}],
    "TNNT2": [{"disease": "Dilated cardiomyopathy", "inheritance": "AD"}, {"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}],
    "TP53": [{"disease": "Li-Fraumeni syndrome", "inheritance": "AD"}],
    "TPM1": [{"disease": "Hypertrophic cardiomyopathy", "inheritance": "AD"}],
    "TRDN": [{"disease": "Catecholaminergic polymorphic ventricular tachycardia", "inheritance": "AR"}, {"disease": "Long QT syndrome", "inheritance": "AR"}],
    "TSC1": [{"disease": "Tuberous sclerosis complex", "inheritance": "AD"}],
    "TSC2": [{"disease": "Tuberous sclerosis complex", "inheritance": "AD"}],
    "TTN": [{"disease": "Dilated cardiomyopathy (truncating variants only)", "inheritance": "AD"}],
    "TTR": [{"disease": "Hereditary transthyretin-related amyloidosis", "inheritance": "AD"}],
    "VHL": [{"disease": "Von Hippel-Lindau syndrome", "inheritance": "AD"}],
    "WT1": [{"disease": "WT1-related Wilms tumor", "inheritance": "AD"}],
}

def _health_variant_gene(v: dict, edits: dict) -> str:
    return _snv_tx_field(v, edits, "gene_symbol") or v.get("gene_symbol") or v.get("GENE") or ""


def _health_variant_hgvs(v: dict, edits: dict) -> str:
    hgvs_c = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_C", "hgvs_c"))
    hgvs_p = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_P", "hgvs_p"))
    return hgvs_c + (f" ({hgvs_p})" if hgvs_p else "")


def _is_health_clinvar_plp(v: dict) -> bool:
    return _clinvar_label(v).lower().replace("_", " ") in {
        "pathogenic",
        "likely pathogenic",
        "pathogenic/likely pathogenic",
        "likely pathogenic/pathogenic",
    }


def _health_selected_ids(report: dict, category: str, candidate_ids: list[str], variants: dict) -> list[str]:
    section = (report.get("secondary_findings") or {}).get(category) or {}
    selected = {str(x) for x in section.get("selected") or []}
    dismissed = {str(x) for x in section.get("dismissed") or []}
    legacy = report.get("panels") or {}
    out = []
    for vid in candidate_ids:
        if vid in dismissed or (legacy.get(vid, {}) or {}).get(category) == "0":
            continue
        if vid in selected or (legacy.get(vid, {}) or {}).get(category) == "V" or _is_health_clinvar_plp(variants.get(vid, {})):
            out.append(vid)
    return out


def _health_acmg_display_order_ids(ids: list[str], variants: dict, edits: dict) -> list[str]:
    ordered: list[str] = []
    used: set[str] = set()
    for group in _ACMG_SF_GROUPS:
        group_genes = set(group["genes"])
        group_ids = [
            vid for vid in ids
            if _health_variant_gene(variants.get(vid, {}), edits.get(vid) or {}) in group_genes
        ]
        ids_by_gene: dict[str, list[str]] = {}
        for vid in group_ids:
            gene = _health_variant_gene(variants.get(vid, {}), edits.get(vid) or {})
            ids_by_gene.setdefault(gene, []).append(vid)
        for gene_ids in ids_by_gene.values():
            for vid in gene_ids:
                if vid not in used:
                    used.add(vid)
                    ordered.append(vid)
    for vid in sorted(set(ids) - used):
        ordered.append(vid)
    return ordered


def _health_acmg_categorized_ids(
    ids: list[str],
    variants: dict,
    edits: dict,
) -> tuple[list[str], list[str]]:
    """Split reportable ACMG findings into disease-risk and carrier groups."""
    ordered_ids = _health_acmg_display_order_ids(ids, variants, edits)
    ids_by_gene: dict[str, list[str]] = {}
    for vid in ordered_ids:
        gene = _health_variant_gene(variants.get(vid, {}), edits.get(vid) or {}) or "?"
        ids_by_gene.setdefault(gene, []).append(vid)

    risk_ids: list[str] = []
    carrier_ids: list[str] = []
    for gene, gene_ids in ids_by_gene.items():
        inheritance_codes = set(_health_acmg_codes(gene))
        is_single_heterozygous_ar = (
            inheritance_codes == {"AR"}
            and len(gene_ids) == 1
            and _health_zygosity_key(
                (variants.get(gene_ids[0]) or {}).get("zygosity", "")
            ) == "het"
        )
        if is_single_heterozygous_ar:
            carrier_ids.extend(gene_ids)
        else:
            risk_ids.extend(gene_ids)
    return risk_ids, carrier_ids


def _acmg_disease_text(gene: str) -> str:
    rows = _ACMG_SF_DISEASES.get(gene) or []
    parts: list[str] = []
    for row in rows:
        disease = row.get("disease") or ""
        if disease:
            parts.append(disease)
    return "；".join(parts)


def _acmg_inheritance_text(gene: str) -> str:
    rows = _ACMG_SF_DISEASES.get(gene) or []
    seen: set[str] = set()
    out: list[str] = []
    for row in rows:
        inh = row.get("inheritance") or ""
        zh = _inheritance_zh(inh)
        if zh and zh not in seen:
            seen.add(zh)
            out.append(zh)
    return "、".join(out)


def _health_acmg_codes(gene: str) -> list[str]:
    return list(dict.fromkeys(
        str(row.get("inheritance") or "").strip()
        for row in (_ACMG_SF_DISEASES.get(gene) or [])
        if str(row.get("inheritance") or "").strip()
    ))


def _health_pathogenicity_zh(label: str) -> str:
    if any(token in str(label or "") for token in ("致病性", "不確定意義", "疑似良性", "良性")):
        return str(label)
    normalized = str(label or "").strip().lower().replace("_", " ")
    if normalized == "pathogenic":
        return "致病性"
    if normalized == "likely pathogenic":
        return "疑似致病性"
    if normalized in {"uncertain significance", "vus"}:
        return "不確定意義"
    if normalized == "likely benign":
        return "疑似良性"
    if normalized == "benign":
        return "良性"
    return f"{label}分級" if label else "致病或疑似致病"


def _health_pathogenicities_zh(labels: Iterable[str]) -> str:
    values = list(dict.fromkeys(
        _health_pathogenicity_zh(label)
        for label in labels
        if str(label or "").strip()
    ))
    return "及".join(values) if values else "致病或疑似致病"


def _health_zygosity_key(value: str) -> str:
    text = str(value or "").strip().lower()
    if "hemi" in text:
        return "hemi"
    if "homo" in text or text == "hom":
        return "hom"
    if "hetero" in text or text == "het":
        return "het"
    return "unknown"


def _health_sex_karyotype(sample_id: str, meta: dict) -> str:
    """Return XX/XY from a DRAGEN ploidy sidecar, then fall back to EMR sex."""
    sample_dir = TERTIARY_OUTPUT_ROOT / sample_id
    karyotype = ploidy.load_sample_ploidy(sample_dir).get("karyotype") or ""
    if karyotype:
        return karyotype
    sex = str(meta.get("Sex") or meta.get("sex") or "").strip().upper()
    if sex in {"M", "MALE", "男"}:
        return "XY"
    if sex in {"F", "FEMALE", "女"}:
        return "XX"
    return ""


def _health_acmg_narrative(
    gene: str,
    *,
    disease_text: str,
    inheritance_codes: list[str],
    acmg: str,
    zygosity: str,
    same_gene_count: int,
    sex_karyotype: str,
) -> list[str]:
    disease = disease_text or "相關遺傳性疾病"
    patho = _health_pathogenicity_zh(acmg)
    zyg = _health_zygosity_key(zygosity)
    codes = set(inheritance_codes)
    inheritance = "、".join(_inheritance_zh(code) for code in inheritance_codes if code)
    standard_advice = (
        "建議至遺傳諮詢或門診相關專科，結合個人病史、家族史及適當的臨床檢查進一步評估。"
        "必要時可考慮對具血緣關係的家屬進行此特定位點之驗證檢測。"
    )
    variable_risk = (
        "本結果表示受檢者可能具有較高的相關疾病風險，但實際是否發病、發病年齡及疾病嚴重程度"
        "可能因個人、家族及環境等因素而異。"
    )

    if codes == {"AD"}:
        return [
            f"{gene} 基因與「{disease}」相關，其遺傳模式為體染色體顯性遺傳。",
            f"此為{patho}之變異位點，{variable_risk}",
            standard_advice,
        ]
    if codes == {"AR"}:
        if zyg == "hom":
            return [
                f"{gene} 基因與「{disease}」相關，其遺傳模式為體染色體隱性遺傳。",
                f"此為{patho}之變異位點，{variable_risk}",
                standard_advice,
            ]
        if same_gene_count >= 2:
            return [
                f"{gene} 基因與「{disease}」相關，其遺傳模式為體染色體隱性遺傳。",
                f"此為{patho}之變異位點。由於兩變異之相位尚未確認，建議進行家族成員檢測，"
                "以釐清兩變異是否位於不同等位基因，方能判斷是否符合體染色體隱性疾病之雙等位基因"
                "致病型態，並評估與相關疾病之關聯性。實際是否發病、發病年齡及疾病嚴重程度亦可能因個人、"
                "家族及環境等因素而異。",
                "建議至遺傳諮詢或門診相關專科，結合個人病史、家族史及適當的臨床檢查進一步評估，"
                "並進行相位分析以釐清變異相位。必要時可考慮對具血緣關係的家屬進行此特定位點之驗證檢測。",
            ]
        return [
            f"{gene} 基因與「{disease}」相關，其遺傳模式為體染色體隱性遺傳。"
            "本次僅檢出一個符合報告條件之變異，檢測結果符合帶因者狀態。",
            f"此為{patho}之變異位點，本結果表示受檢者可能為相關疾病的帶因者。本檢測僅針對"
            "疾病資料庫中已收錄之致病性或疑似致病性變異，故仍可能存在資料庫尚未收錄或本檢測方法未涵蓋"
            "的其他變異。",
            standard_advice,
        ]
    if codes and codes.issubset({"XL", "XLR", "XLD"}):
        if sex_karyotype == "XY":
            return [
                f"{gene} 基因與「{disease}」相關，其遺傳模式為性聯遺傳。",
                f"此為{patho}之變異位點，{variable_risk}",
                standard_advice,
            ]
        if sex_karyotype == "XX":
            return [
                f"{gene} 基因與「{disease}」相關，其遺傳模式為性聯遺傳。",
                f"此為{patho}之變異位點。女性帶有單一變異時，臨床表現可能受 X 染色體失活型態"
                "及疾病本身表現範圍影響，可能無症狀或出現不同程度之相關表徵。"
                "實際是否發病、發病年齡及疾病嚴重程度亦可能因個人、家族及環境等因素而異。",
                standard_advice,
            ]
        return [
            f"{gene} 基因與「{disease}」相關，其遺傳模式為性聯遺傳。",
            f"此為{patho}之變異位點；目前無足夠的性染色體組成資訊，無法進一步區分其疾病風險，"
            "仍需結合臨床資料綜合評估。",
            standard_advice,
        ]
    return [
        f"{gene} 基因與「{disease}」相關{f'，其遺傳模式為{inheritance}' if inheritance else ''}。",
        f"此為{patho}之變異位點，{variable_risk}",
        standard_advice,
    ]


def _health_snv_variant_block(doc, v: dict, *, tier: str, edits: dict,
                              disease_text: str = "",
                              inheritance_text: str = "",
                              inheritance_codes: list[str] | None = None,
                              same_gene_count: int = 1,
                              sex_karyotype: str = "") -> None:
    gene = _health_variant_gene(v, edits) or "?"
    tx = _snv_transcript_label(v, edits)
    _add_paragraph(doc, f"    {gene} ({tx})", bold=True)
    rs = v.get("rs_id") or v.get("RS_ID") or ""
    struc = _structure_label({
        **v,
        "exon": _snv_tx_field(v, edits, "exon"),
        "intron": _snv_tx_field(v, edits, "intron"),
    })
    hgvs_c = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_C", "hgvs_c"))
    hgvs_p = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_P", "hgvs_p"))
    nuc = hgvs_c + (f"({hgvs_p})" if hgvs_p else "")
    zyg = _zygosity_long(v.get("zygosity", ""))
    clnsig = _clinvar_label(v)
    acmg_raw = _acmg_label(v, edits)

    _ascii_table(doc, columns=[
        ("基因",          9),
        ("RS ID",         7, "buffered"),
        ("結構",          9),
        ("核苷酸",       22, "hgvs"),
        ("基因型",       13),
        ("ClinVar",      16, "token-buffered"),
        ("ACMG&AMP指引", 12, "token"),
    ], rows=[[gene, rs, struc, nuc, zyg, clnsig, acmg_raw]])

    if disease_text and inheritance_codes:
        narrative = _health_acmg_narrative(
            gene,
            disease_text=disease_text,
            inheritance_codes=inheritance_codes,
            acmg=acmg_raw,
            zygosity=zyg,
            same_gene_count=same_gene_count,
            sex_karyotype=sex_karyotype,
        )
        for idx, text in enumerate(narrative, start=1):
            _add_paragraph(doc, f"    {idx}. {text}")
    else:
        _add_paragraph(doc, f"    1. {_omim_block_for_snv(v, edits)}")
        _add_paragraph(doc, f"    2. {_patho_sentence(acmg_raw)}")


def _health_snv_gene_block(
    doc,
    rows: list[tuple[dict, dict]],
    *,
    disease_text: str,
    inheritance_codes: list[str],
    sex_karyotype: str,
) -> None:
    first_v, first_edits = rows[0]
    gene = _health_variant_gene(first_v, first_edits) or "?"
    tx_labels = list(dict.fromkeys(
        _snv_transcript_label(v, edits)
        for v, edits in rows
        if _snv_transcript_label(v, edits)
    ))
    heading = f"    {gene}"
    if tx_labels:
        heading += f" ({'; '.join(tx_labels)})"
    _add_paragraph(doc, heading, bold=True)

    table_rows = []
    acmg_labels: list[str] = []
    zygosities: list[str] = []
    for v, edits in rows:
        struc = _structure_label({
            **v,
            "exon": _snv_tx_field(v, edits, "exon"),
            "intron": _snv_tx_field(v, edits, "intron"),
        })
        hgvs_c = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_C", "hgvs_c"))
        hgvs_p = _strip_hgvs_prefix(_snv_tx_field(v, edits, "HGVS_P", "hgvs_p"))
        zyg = _zygosity_long(v.get("zygosity", ""))
        acmg_raw = _acmg_label(v, edits)
        acmg_labels.append(acmg_raw)
        zygosities.append(zyg)
        table_rows.append([
            gene,
            v.get("rs_id") or v.get("RS_ID") or "",
            struc,
            hgvs_c + (f"({hgvs_p})" if hgvs_p else ""),
            zyg,
            _clinvar_label(v),
            acmg_raw,
        ])

    _ascii_table(doc, columns=[
        ("基因",          9),
        ("RS ID",         7, "buffered"),
        ("結構",          9),
        ("核苷酸",       20, "hgvs"),
        ("基因型",       14),
        ("ClinVar",      16, "token-buffered"),
        ("ACMG&AMP指引", 12, "token"),
    ], rows=table_rows)

    narrative = _health_acmg_narrative(
        gene,
        disease_text=disease_text,
        inheritance_codes=inheritance_codes,
        acmg=_health_pathogenicities_zh(acmg_labels),
        zygosity=zygosities[0] if len(set(zygosities)) == 1 else "",
        same_gene_count=len(rows),
        sex_karyotype=sex_karyotype,
    )
    for idx, text in enumerate(narrative, start=1):
        _add_paragraph(doc, f"    {idx}. {text}")


def _render_health_secondary_section(doc, title: str, ids: list[str], variants: dict, report: dict) -> None:
    _add_paragraph(doc, title, bold=True)
    if not ids:
        _add_paragraph(doc, f"  {_NO_HEALTH_VARIANT_TEXT}")
    else:
        edits = report.get("edits") or {}
        for idx, vid in enumerate(ids, start=1):
            v = variants.get(vid) or {}
            _health_snv_variant_block(doc, v, tier=str(idx), edits=edits.get(vid) or {})
            _blank(doc)
    _blank(doc)


def _render_health_acmg_section(
    doc,
    title: str,
    ids: list[str],
    variants: dict,
    report: dict,
    *,
    sex_karyotype: str = "",
) -> None:
    _add_paragraph(doc, f"  {_HEALTH_ACMG_CAUTION}")
    _blank(doc)
    _add_paragraph(doc, title, bold=True)
    edits = report.get("edits") or {}
    risk_ids, carrier_ids = _health_acmg_categorized_ids(ids, variants, edits)

    if not risk_ids:
        _add_paragraph(doc, f"  {_NO_HEALTH_VARIANT_TEXT}")
    else:
        ids_by_gene: dict[str, list[str]] = {}
        for vid in risk_ids:
            gene = _health_variant_gene(variants.get(vid, {}), edits.get(vid) or {}) or "?"
            ids_by_gene.setdefault(gene, []).append(vid)
        gene_groups = list(ids_by_gene.items())
        for group_index, (gene, gene_ids) in enumerate(gene_groups):
            _health_snv_gene_block(
                doc,
                [(variants.get(vid) or {}, edits.get(vid) or {}) for vid in gene_ids],
                disease_text=_acmg_disease_text(gene),
                inheritance_codes=_health_acmg_codes(gene),
                sex_karyotype=sex_karyotype,
            )
            if group_index < len(gene_groups) - 1:
                _blank(doc)

    _blank(doc)
    _add_paragraph(
        doc,
        "第二類：符合帶因者狀態之致病性或疑似致病性變異位點",
        bold=True,
    )
    if not carrier_ids:
        _add_paragraph(doc, f"  {_NO_HEALTH_CARRIER_VARIANT_TEXT}")
    else:
        ids_by_gene = {}
        for vid in carrier_ids:
            gene = _health_variant_gene(variants.get(vid, {}), edits.get(vid) or {}) or "?"
            ids_by_gene.setdefault(gene, []).append(vid)
        for gene, gene_ids in ids_by_gene.items():
            _health_snv_gene_block(
                doc,
                [(variants.get(vid) or {}, edits.get(vid) or {}) for vid in gene_ids],
                disease_text=_acmg_disease_text(gene),
                inheritance_codes=_health_acmg_codes(gene),
                sex_karyotype=sex_karyotype,
            )
            _blank(doc)

    _add_paragraph(doc, f"  {_HEALTH_NEGATIVE_LIMITATION}")
    _blank(doc)


def _pgx_no_action(text: str) -> bool:
    value = str(text or "").strip().lower()
    return any(phrase in value for phrase in (
        "no action is required",
        "no action is needed",
        "no recommendation",
        "no need to avoid",
        "per standard dosing",
        "standard prescribing",
    ))


def _pgx_annotation_genes(annotation: dict) -> list[str]:
    return [
        str(gene).strip()
        for gene in annotation.get("genes") or []
        if str(gene).strip() in _PGX_CPIC_LEVEL_A_SET
    ]


def _pgx_annotation_is_actionable(annotation: dict) -> bool:
    section = str(annotation.get("section") or "")
    classification = str(annotation.get("classification") or "").strip().lower()
    recommendation = str(annotation.get("recommendation") or "")
    if not recommendation or _pgx_no_action(recommendation):
        return False
    has_action = any((
        annotation.get("dosing_information"),
        annotation.get("alternate_drug_available"),
        annotation.get("other_prescribing_guidance"),
    ))
    if section == "CPIC Guideline Annotation":
        return classification in {"strong", "moderate"} and has_action
    if section == "FDA PGx Association":
        return annotation.get("fda_category") == "therapeutic_management"
    if section == "FDA Label Annotation":
        return has_action
    return False


def _pgx_summary_alerts(pgx: dict) -> list[dict]:
    best_by_gene_drug: dict[tuple[str, str], dict] = {}
    seen: set[tuple[str, str, str, str]] = set()
    for item in pgx.get("guideline_annotations") or []:
        section = str(item.get("section") or "")
        classification = str(item.get("classification") or "").strip()
        recommendation = item.get("recommendation") or ""
        if not _pgx_annotation_is_actionable(item):
            continue
        if section == "FDA Label Annotation":
            continue
        is_cpic = section == "CPIC Guideline Annotation"
        source = "CPIC" if is_cpic else "FDA"
        level = classification if is_cpic else "Therapeutic management"
        rank = 0 if classification.lower() == "strong" else (1 if is_cpic else 2)
        for gene in _pgx_annotation_genes(item):
            if not _pgx_gene_has_actionable_result(pgx, gene):
                continue
            _, phenotype = _pgx_gene_result(pgx, gene)
            if any(token in phenotype.lower() for token in ("indeterminate", "uncertain susceptibility")):
                continue
            key = (gene, item.get("drug") or "", source, recommendation)
            if key in seen:
                continue
            seen.add(key)
            alert = {
                "gene": gene,
                "drug": item.get("drug") or "",
                "source": source,
                "level": level,
                "recommendation": recommendation,
                "rank": rank,
            }
            gene_drug = (gene, str(item.get("drug") or "").lower())
            current = best_by_gene_drug.get(gene_drug)
            if current is None or alert["rank"] < current["rank"]:
                best_by_gene_drug[gene_drug] = alert
    alerts = list(best_by_gene_drug.values())
    alerts.sort(key=lambda row: (row["rank"], row["gene"], row["drug"]))
    return alerts


def _pgx_gene_result(pgx: dict, gene: str) -> tuple[str, str]:
    genes = pgx.get("genes") or {}
    payload = genes.get(gene) or {}
    details = payload.get("details") or {}
    diplotype = details.get("label") or ""
    phenotype = "；".join(details.get("phenotypes") or [])
    raw_phenotype = phenotype
    tsv_diplotype = payload.get("diplotype") or ""
    tsv_phenotype = payload.get("phenotype") or ""
    if (not diplotype or (gene == "MT-RNR1" and diplotype.lower() == "unknown")) and tsv_diplotype:
        diplotype = tsv_diplotype
    if (
        not phenotype
        or (gene == "MT-RNR1" and phenotype.lower() in {"no result", "unknown"})
    ) and tsv_phenotype:
        phenotype = tsv_phenotype
        raw_phenotype = phenotype
    allele_function = _pgx_allele_function(details)
    diplotype, _ = _pgx_display_genotype(gene, diplotype, raw_phenotype)
    if gene == "VKORC1" and allele_function:
        phenotype = allele_function
    elif not _pgx_display_phenotype(phenotype) and allele_function:
        phenotype = allele_function
    return diplotype, phenotype


def _pgx_gene_has_actionable_result(pgx: dict, gene: str) -> bool:
    """Return whether this gene, rather than a co-annotated gene, drives action."""
    _, phenotype = _pgx_gene_result(pgx, gene)
    normalized = str(phenotype or "").strip().lower()
    if not normalized:
        return True
    if normalized in {"normal", "negative", "low risk"}:
        return False
    return not any(token in normalized for token in (
        "normal metabolizer",
        "normal function",
        "normal activity",
        "normal risk",
        "no increased risk",
        "negative",
        "low risk",
        "indeterminate",
        "uncertain susceptibility",
    ))


def _pgx_display_phenotype(phenotype: str) -> str:
    value = str(phenotype or "").strip()
    return "" if value.lower() in {"", "-", "—", "n/a", "na"} else phenotype


def _pgx_allele_function(details: dict) -> str:
    values: list[str] = []
    for key in ("allele1_function", "allele2_function"):
        value = str(details.get(key) or "").strip()
        if value and _pgx_display_phenotype(value) and value not in values:
            values.append(value)
    return "；".join(values)


def _pgx_genotype_rows(pgx: dict) -> list[list[str]]:
    rows: list[list[str]] = []
    for gene in _PGX_CPIC_LEVEL_A_GENES:
        diplotype, phenotype = _pgx_gene_result(pgx, gene)
        rows.append([gene, diplotype or "—", _pgx_display_phenotype(phenotype) or "No phenotype assigned"])
    return rows


def _pgx_source_label(section: str, annotation: dict) -> tuple[str, str, int]:
    classification = str(annotation.get("classification") or "").strip()
    if section == "CPIC Guideline Annotation":
        level = classification
        priority = 0 if classification.lower() == "strong" else 1
        return "CPIC", level, priority
    if section == "FDA Label Annotation":
        return "FDA Label", classification or "Unspecified", 3
    if section == "FDA PGx Association":
        category = annotation.get("fda_category") or "unspecified"
        return "FDA PGx Association", category.replace("_", " ").title(), 2
    return section, classification, 9


def _pgx_recommendation_key(text: str) -> str:
    value = str(text or "").lower()
    value = re.sub(r"see label for more information\.?", "", value)
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _pgx_source_level_display(source: str, level: str) -> str:
    if source == "FDA PGx Association":
        return "FDA Therapeutic Management"
    if source == "FDA Label":
        return "FDA Label"
    return f"{source} {level}".strip()


def _pgx_drug_groups(pgx: dict) -> list[dict]:
    grouped: dict[str, dict] = {}
    for annotation in pgx.get("guideline_annotations") or []:
        section = str(annotation.get("section") or "")
        if section not in {
            "CPIC Guideline Annotation",
            "FDA Label Annotation",
            "FDA PGx Association",
        }:
            continue
        recommendation = str(annotation.get("recommendation") or "")
        if not _pgx_annotation_is_actionable(annotation):
            continue
        drug = _pgx_drug_label(annotation.get("drug") or "")
        if not drug:
            continue
        source, level, source_priority = _pgx_source_label(section, annotation)
        action = _pgx_action_zh(recommendation)
        group = grouped.setdefault(drug.casefold(), {
            "drug": drug,
            "genes": {},
            "recommendations": [],
        })
        for gene in _pgx_annotation_genes(annotation):
            if not _pgx_gene_has_actionable_result(pgx, gene):
                continue
            diplotype, phenotype = _pgx_gene_result(pgx, gene)
            if any(token in phenotype.lower() for token in ("indeterminate", "uncertain susceptibility")):
                continue
            group["genes"][gene] = {
                "gene": gene,
                "diplotype": diplotype,
                "phenotype": phenotype,
            }
            matching = next((
                row for row in group["recommendations"]
                if row.get("gene") == gene
                and row.get("source") == source
                and row.get("action") == action
            ), None)
            if matching:
                if source == "CPIC" and source_priority < matching.get("source_priority", 9):
                    matching.update({
                        "level": level,
                        "recommendation": recommendation,
                        "source_priority": source_priority,
                        "rec_key": _pgx_recommendation_key(recommendation),
                    })
                elif len(recommendation) > len(str(matching.get("recommendation") or "")):
                    matching["recommendation"] = recommendation
                    matching["rec_key"] = _pgx_recommendation_key(recommendation)
                continue
            rec_key = _pgx_recommendation_key(recommendation)
            if any(
                row.get("gene") == gene
                and row.get("source") == source
                and row.get("level") == level
                and row.get("rec_key") == rec_key
                for row in group["recommendations"]
            ):
                continue
            group["recommendations"].append({
                "gene": gene,
                "source": source,
                "level": level,
                "recommendation": recommendation,
                "action": action,
                "source_priority": source_priority,
                "rec_key": rec_key,
            })
    out: list[dict] = []
    for group in grouped.values():
        if not group["recommendations"]:
            continue
        group["recommendations"].sort(key=lambda row: (
            row.get("source_priority", 9),
            row.get("gene") or "",
            row.get("source") or "",
            row.get("recommendation") or "",
        ))
        action_rank = {
            "調整劑量並監測": 0,
            "考慮替代藥物": 1,
            "加強不良反應監測": 2,
            "使用前確認表型或檢驗": 3,
            "參考最新藥品仿單": 4,
        }
        primary = min(
            group["recommendations"],
            key=lambda row: (action_rank.get(row.get("action") or "", 9), row.get("source_priority", 9)),
        )
        group["action"] = primary.get("action") or "參考最新藥品仿單"
        group["source_level"] = _pgx_source_level_display(
            primary.get("source") or "",
            primary.get("level") or "",
        )
        out.append(group)
    return sorted(out, key=lambda group: str(group.get("drug") or "").casefold())


def _pgx_full_groups(pgx: dict) -> list[dict]:
    grouped: dict[str, dict] = {}
    for annotation in pgx.get("guideline_annotations") or []:
        section = str(annotation.get("section") or "")
        if section not in {
            "CPIC Guideline Annotation",
            "FDA Label Annotation",
            "FDA PGx Association",
        }:
            continue
        recommendation = annotation.get("recommendation") or ""
        if not _pgx_annotation_is_actionable(annotation):
            continue
        for gene in _pgx_annotation_genes(annotation):
            if not _pgx_gene_has_actionable_result(pgx, gene):
                continue
            diplotype, phenotype = _pgx_gene_result(pgx, gene)
            if any(token in phenotype.lower() for token in ("indeterminate", "uncertain susceptibility")):
                continue
            group = grouped.setdefault(gene, {
                "gene": gene,
                "diplotype": diplotype,
                "phenotype": phenotype,
                "recommendations": [],
            })
            source = "CPIC" if section.startswith("CPIC") else (
                "FDA Label" if section == "FDA Label Annotation" else "FDA PGx Association"
            )
            classification = str(annotation.get("classification") or "").strip()
            if section == "FDA PGx Association":
                category = annotation.get("fda_category") or "unspecified"
                level = category.replace("_", " ").title()
            else:
                level = classification
            matching = next((
                row for row in group["recommendations"]
                if (
                    row.get("drug") or "",
                    row.get("source") or "",
                    row.get("recommendation") or "",
                ) == (annotation.get("drug") or "", source, recommendation)
            ), None)
            if matching:
                strength_rank = {"strong": 0, "moderate": 1, "optional": 2}
                if strength_rank.get(level.lower(), 9) < strength_rank.get(
                    str(matching.get("level") or "").lower(), 9
                ):
                    matching["level"] = level
                continue
            group["recommendations"].append({
                "drug": annotation.get("drug") or "",
                "source": source,
                "recommendation": recommendation,
                "level": level,
            })
    for group in grouped.values():
        group["recommendations"].sort(key=lambda row: (
            0 if row["source"] == "CPIC" else 1,
            row["drug"],
            row["source"],
        ))
    return sorted(grouped.values(), key=lambda group: group["gene"])


def _pgx_display_genotype(gene: str, diplotype: str, phenotype: str) -> tuple[str, str]:
    if gene != "VKORC1":
        return diplotype, phenotype
    alleles = re.findall(r"rs9923231 variant \(([A-Z])\)", diplotype or "", flags=re.I)
    normalized_pheno = re.sub(r"^-1639\s*([A-Z])([A-Z])$", r"-1639 \1/\2", phenotype or "")
    if len(alleles) == 2:
        genotype = f"rs9923231 {alleles[0].upper()}/{alleles[1].upper()}"
        if normalized_pheno:
            genotype += f"（{normalized_pheno}）"
        return genotype, "—"
    return diplotype, phenotype


def _pgx_drug_label(drug: str) -> str:
    text = str(drug or "").strip()
    return text[:1].upper() + text[1:] if text else ""


def _pgx_action_zh(recommendation: str) -> str:
    text = str(recommendation or "").lower()
    if any(token in text for token in (
        "avoid", "alternative", "contraindicated", "do not use", "not recommended",
    )):
        return "考慮替代藥物"
    if any(token in text for token in ("dose", "dosage", "dosing", "titrate", "reduction", "increase")):
        return "調整劑量並監測"
    if any(token in text for token in ("monitor", "adverse", "caution", "toxicity", "risk")):
        return "加強不良反應監測"
    if any(token in text for token in ("confirm", "test", "phenotype", "activity")):
        return "使用前確認表型或檢驗"
    return "參考最新藥品仿單"


def _pgx_summary_rows(alerts: list[dict]) -> list[dict]:
    rows: list[dict] = []
    for alert in alerts:
        action = _pgx_action_zh(alert.get("recommendation") or "")
        rows.append({
            "drug": _pgx_drug_label(alert.get("drug") or ""),
            "gene": alert.get("gene") or "",
            "action": action,
            "source_level": f"{alert.get('source') or ''} {alert.get('level') or ''}".strip(),
        })
    return sorted(rows, key=lambda row: (
        str(row.get("drug") or "").casefold(),
        str(row.get("gene") or "").casefold(),
        str(row.get("source_level") or "").casefold(),
    ))


def _pgx_summary_rows_from_drug_groups(groups: list[dict]) -> list[dict]:
    rows: list[dict] = []
    action_rank = {
        "調整劑量並監測": 0,
        "考慮替代藥物": 1,
        "加強不良反應監測": 2,
        "使用前確認表型或檢驗": 3,
        "參考最新藥品仿單": 4,
    }
    for group in sorted(groups, key=lambda row: str(row.get("drug") or "").casefold()):
        summary_recs = [
            rec for rec in group.get("recommendations") or []
            if rec.get("source") in {"CPIC", "FDA PGx Association"}
        ]
        if not summary_recs:
            continue
        primary = min(
            summary_recs,
            key=lambda rec: (
                action_rank.get(rec.get("action") or "", 9),
                rec.get("source_priority", 9),
            ),
        )
        genes = [rec.get("gene") for rec in summary_recs if rec.get("gene")]
        rows.append({
            "drug": group["drug"],
            "gene": "、".join(dict.fromkeys(genes)),
            "action": primary.get("action") or "",
            "source_level": _pgx_source_level_display(
                primary.get("source") or "",
                primary.get("level") or "",
            ),
        })
    return rows


def _pgx_action_categories(groups: list[dict]) -> list[tuple[str, list[str]]]:
    labels = [
        "調整劑量並監測",
        "考慮替代藥物",
        "加強不良反應監測",
        "使用前確認表型或檢驗",
        "參考最新藥品仿單",
    ]
    by_action: dict[str, list[str]] = {label: [] for label in labels}
    for group in groups:
        action = group.get("action") or "參考最新藥品仿單"
        by_action.setdefault(action, []).append(group["drug"])
    return [
        (label, sorted(dict.fromkeys(drugs), key=str.casefold))
        for label, drugs in by_action.items()
        if drugs
    ]


def _pgx_gene_phenotype_text(group: dict) -> str:
    parts = []
    for gene, payload in sorted(group.get("genes", {}).items()):
        phenotype = payload.get("phenotype") or "—"
        parts.append(f"{gene} {phenotype}")
    return "；".join(parts)


def _pgx_clean_recommendation_text(text: str) -> str:
    value = html.unescape(str(text or "").strip())
    value = value.replace('"', "")
    value = value.replace("“", "").replace("”", "")
    value = re.sub(r"\[\s*\.\.\.\s*\]", "", value)
    value = re.sub(r"(?<=[a-z0-9])\.\.\.(?=[A-Z])", ". ", value)
    value = value.replace("...", " ")
    value = re.sub(r"(?<![A-Za-z0-9])\.\.\.(?![A-Za-z0-9])", "", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _pgx_recommendation_text(group: dict) -> str:
    cpic_parts = []
    fda_therapeutic_texts: list[str] = []
    fda_label_texts: list[str] = []
    for rec in group.get("recommendations") or []:
        source = rec.get("source") or ""
        level = rec.get("level") or ""
        recommendation = _pgx_clean_recommendation_text(rec.get("recommendation") or "")
        if source == "CPIC":
            label = _pgx_source_level_display(source, level)
            cpic_parts.append(f"{recommendation} ({label})" if label else recommendation)
            continue
        if source == "FDA PGx Association":
            fda_therapeutic_texts.append(recommendation)
            continue
        if source == "FDA Label":
            fda_label_texts.append(recommendation)
            continue
        label = _pgx_source_level_display(source, level)
        cpic_parts.append(f"{recommendation} ({label})" if label else recommendation)
    parts = cpic_parts[:]
    fda_texts = fda_therapeutic_texts + fda_label_texts
    if fda_texts:
        fda_sources = []
        if fda_therapeutic_texts:
            fda_sources.append("FDA Therapeutic Management")
        if fda_label_texts:
            fda_sources.append("FDA Label")
        parts.append(f"{' '.join(fda_texts)} ({' / '.join(fda_sources)})")
    return "；".join(parts)


def _render_health_pgx_resources(doc) -> None:
    _blank(doc)
    _add_paragraph(doc, "  官方用藥資訊查詢", bold=True)
    _add_paragraph(doc, f"  {_HEALTH_PGX_RESOURCE_INTRO}")
    for index, (label, url) in enumerate(_HEALTH_PGX_RESOURCES, start=1):
        paragraph = _add_paragraph(doc, f"    {index}. {label}：")
        _add_hyperlink(paragraph, url, url)
    _add_paragraph(doc, f"  {_HEALTH_PGX_RESOURCE_CLOSING}")
    _blank(doc)


def _render_health_pgx_section(doc, title: str, pgx: dict) -> list[dict]:
    _add_paragraph(doc, title, bold=True)
    _add_paragraph(doc, f"  {_HEALTH_PGX_CAUTION}")
    _blank(doc)
    drug_groups = _pgx_drug_groups(pgx or {})
    if not drug_groups:
        _add_paragraph(doc, "  本次檢測未發現符合目前回報規則之明確臨床可應用藥物基因體結果。")
        _add_paragraph(
            doc,
            "  此結果不代表所有藥物均不會發生療效差異或不良反應。藥物反應亦會受到年齡、"
            "疾病狀態、肝腎功能、併用藥物、生活習慣及其他基因因素影響。",
        )
    else:
        genes = "、".join(dict.fromkeys(
            gene
            for group in drug_groups
            for gene in group.get("genes", {})
        ))
        summary_rows = _pgx_summary_rows_from_drug_groups(drug_groups)
        summary_drugs = list(dict.fromkeys(row["drug"] for row in summary_rows if row.get("drug")))
        drug_text = "、".join(summary_drugs[:12])
        _add_paragraph(
            doc,
            f"  本次檢測發現 {len(drug_groups)} 項具臨床用藥參考價值的藥物結果"
            f"{f'，涉及 {genes} 基因' if genes else ''}"
            f"{f'，可能影響 {drug_text} 及其他藥物之使用' if drug_text else ''}。"
            "此處列出用藥建議概覽與摘要，完整藥物建議詳見報告末端附錄。"
            "若目前使用或未來考慮使用"
            "相關藥物，建議由處方醫師參考下方結果或最新 FDA/CPIC 指引進行評估。",
        )
        _blank(doc)
        _add_paragraph(doc, "  用藥建議概覽", bold=True)
        for action, drugs in _pgx_action_categories(drug_groups):
            _add_paragraph(doc, f"    {action}：{'、'.join(drugs)}")
        _add_paragraph(
            doc,
            "    其餘未列之藥物，未發現符合本報告回報規則之明確處方調整建議。",
        )
        _blank(doc)
        _add_paragraph(doc, "  藥物建議摘要", bold=True)
        _ascii_table(doc, columns=[
            ("藥物", 24, "word-buffered"),
            ("基因", 14, "buffered"),
            ("建議處置", 18, "buffered"),
            ("建議依據及等級", 27, "buffered"),
        ], rows=[
            [row["drug"], row["gene"], row["action"], row["source_level"]]
            for row in summary_rows
        ], indent="  ")
    _blank(doc)
    _add_paragraph(doc, "  基因型與表現型", bold=True)
    _ascii_table(doc, columns=[
        ("基因", 12),
        ("基因型", 30, "genotype-buffered"),
        ("表型", 42, "buffered"),
    ], rows=_pgx_genotype_rows(pgx or {}), indent="  ")
    _render_health_pgx_resources(doc)
    return drug_groups


def _render_health_pgx_appendix(doc, title: str, drug_groups: list[dict]) -> None:
    _add_paragraph(doc, title, bold=True)
    _ascii_table(doc, columns=_PGX_FULL_RECOMMENDATION_COLUMNS, rows=[
        [
            group["drug"],
            _pgx_gene_phenotype_text(group),
            _pgx_recommendation_text(group),
        ]
        for group in drug_groups
    ], indent="  ")


def _health_panel_gene_sections(requested_set: set[str]) -> list[tuple[str, list[str]]]:
    out: list[tuple[str, list[str]]] = []
    panel_keys = {
        "acmg_sf": "ACMG_SF_v3.3",
        "lipid_fh": "lipid_fh",
        "hereditary_cancer": "WES-I__腫瘤醫學__遺傳癌症",
        "stroke": "WGS__神經科__Stroke",
        "carrier": "carrier_mackenzie_1300+",
        "proactive": "proactive",
    }
    titles = dict(_HEALTH_SECTION_ORDER)
    titles["acmg_sf"] = _HEALTH_ACMG_GENE_LIST_TITLE
    for key, panel_name in panel_keys.items():
        if key not in requested_set:
            continue
        payload = phenotype_scorer.genes_for_key(panel_name, kind="panel")
        genes = sorted({
            panel_deadzone.canonical_panel_gene_symbol(g)
            for g in payload.get("genes", [])
            if g
        })
        out.append((titles.get(key, key), [g for g in genes if g]))
    return out


def _health_test_bundle_name(requested_set: set[str]) -> str:
    has_pgx = "pgx" in requested_set
    has_disease = any(section != "pgx" for section in requested_set)
    if has_pgx and has_disease:
        return "ACMG疾病風險基因及藥物基因體學基因篩檢"
    if has_pgx:
        return "藥物基因體學基因篩檢"
    return "ACMG疾病風險基因篩檢"


def _pgx_report_genes(pgx: dict) -> list[str]:
    return list(_PGX_CPIC_LEVEL_A_GENES)


def _section_health_annotations(doc, requested_set: set[str], pgx: dict | None = None) -> None:
    _add_paragraph(doc, "五、檢測結果注釋")
    _add_paragraph(doc, "  1. 本檢測結果比對參考序列為人類hg38版本。")
    _add_paragraph(doc, "  2. 參考資料:")
    _add_paragraph(doc, f"     a. 疾病資料庫: OMIM、ClinVar ({CLINVAR_DATE})")
    _add_paragraph(doc, "     b. 族群資料庫: gnomAD (v4.1 genome)")
    _add_paragraph(doc, "     c. 序列資料庫: RefSeqGene (105.20220307)")
    _add_paragraph(doc, "  3. 本次檢測基因包括")
    sections = _health_panel_gene_sections(requested_set)
    if not sections and "pgx" not in requested_set:
        _add_paragraph(doc, "    （未選擇檢測基因項目）")
        return
    for idx, (name, genes) in enumerate(sections):
        if idx:
            _blank(doc)
        _add_paragraph(doc, f"{name}:")
        if name == _HEALTH_ACMG_GENE_LIST_TITLE:
            for group_index, group in enumerate(_ACMG_SF_GROUPS, start=1):
                _add_paragraph(doc, f"  {_health_acmg_group_label(group, group_index)}")
        elif genes:
            _add_paragraph(doc, ", ".join(genes))
        else:
            _add_paragraph(doc, "依藥物基因體學分析模組可判讀基因輸出。")
    if "pgx" in requested_set:
        if sections:
            _blank(doc)
        _add_paragraph(doc, f"{_HEALTH_PGX_GENE_LIST_TITLE}:")
        genes = _pgx_report_genes(pgx or {})
        if genes:
            _add_paragraph(doc, ", ".join(genes))
        else:
            _add_paragraph(doc, "本次藥物基因體學分析未輸出可列示之基因。")


def _render_health_variant_reference_appendix(
    doc,
    title: str,
    variants: list[dict],
    report: dict,
) -> None:
    _add_paragraph(doc, title, bold=True)
    edits = report.get("edits") or {}
    valid_variants = [variant for variant in variants if variant]
    for index, variant in enumerate(valid_variants):
        _add_paragraph(doc, _snv_reference_text(
            variant,
            edits.get(variant.get("id", ""), {}),
            acmg_zh=True,
        ))
        if index < len(valid_variants) - 1:
            _blank(doc)


def build_health_docx(sample_id: str, *, sections: Iterable[str] | None = None) -> bytes:
    sample = sample_loader.load_sample(sample_id, include_aux=False)
    if sample is None:
        raise FileNotFoundError(f"sample not found: {sample_id}")
    secondary = sample_loader.load_sample_secondary_snv(sample_id) or {}
    pgx_payload = sample_loader.load_sample_pgx(sample_id) or {}

    requested = [str(s).strip() for s in (sections or []) if str(s).strip()]
    if not requested:
        requested = ["acmg_sf", "pgx"]
    requested_set = set(requested)
    report = report_store.load(sample_id)
    variants = secondary.get("variants") or {}
    categories = secondary.get("categories") or {}

    doc = Document()
    _apply_normal_font(doc)
    _apply_page_margins(doc)

    meta = sample.get("meta") or {}
    test_type = meta.get("Test", "") or "WES"

    _section_health_patient_header(doc, meta)
    _section_test_info(doc, test_type, health=True)
    _add_paragraph(doc, f"二、檢驗套組: {_health_test_bundle_name(requested_set)}")
    _blank(doc)

    _add_paragraph(doc, "三、檢測結果")
    _add_paragraph(doc, "  檢體說明:")
    _add_paragraph(doc, "    檢體類別：血液")
    _add_paragraph(doc, "  綜合說明:")

    referenced: list[dict] = []
    referenced_ids: set[str] = set()
    pgx_drug_groups: list[dict] = []
    for key, title in _HEALTH_SECTION_ORDER:
        if key not in requested_set:
            continue
        if key == "pgx":
            pgx_drug_groups = _render_health_pgx_section(
                doc,
                title,
                pgx_payload.get("pgx") or pgx_payload.get("pharmcat") or {},
            )
            continue
        ids = _health_selected_ids(report, key, categories.get(key) or [], variants)
        if key == "acmg_sf":
            risk_ids, carrier_ids = _health_acmg_categorized_ids(
                ids,
                variants,
                report.get("edits") or {},
            )
            for vid in risk_ids + carrier_ids:
                if vid in referenced_ids:
                    continue
                referenced_ids.add(vid)
                referenced.append(variants.get(vid) or {})
            _render_health_acmg_section(
                doc,
                title,
                ids,
                variants,
                report,
                sex_karyotype=_health_sex_karyotype(sample_id, meta),
            )
        else:
            for vid in ids:
                if vid in referenced_ids:
                    continue
                referenced_ids.add(vid)
                referenced.append(variants.get(vid) or {})
            _render_health_secondary_section(doc, title, ids, variants, report)

    _section_methods(doc, test_type, health=True)
    _section_health_annotations(doc, requested_set, pgx_payload.get("pgx") or pgx_payload.get("pharmcat") or {})

    if referenced or pgx_drug_groups:
        _blank(doc)
        _add_paragraph(doc, "附錄", bold=True)
    if referenced:
        _render_health_variant_reference_appendix(
            doc,
            "變異位點參考資料",
            referenced,
            report,
        )
    if pgx_drug_groups:
        if referenced:
            _blank(doc)
        _render_health_pgx_appendix(
            doc,
            "完整用藥建議",
            pgx_drug_groups,
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Top-level entrypoint ──────────────────────────────────────────

def build_diagnosis_docx(sample_id: str, *, gene_list_mode: str = "grouped") -> bytes:
    """Render the full diagnostic report as DOCX bytes.

    gene_list_mode = "grouped" (default) or "merged" controls how §五.4
    "本次檢測基因包括" is rendered.
    """
    sample = sample_loader.load_sample(sample_id, include_aux=True)
    if sample is None:
        raise FileNotFoundError(f"sample not found: {sample_id}")

    report = report_store.load(sample_id)
    meta   = sample.get("meta") or {}
    test_type = meta.get("Test", "") or "WES"

    doc = Document()
    _apply_normal_font(doc)
    _apply_page_margins(doc)

    _section_test_info(doc, test_type)
    _section_panel_set(doc)
    _section_results(doc, sample, report, test_type)
    _section_methods(doc, test_type)
    _section_annotations(doc, sample, gene_list_mode)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
