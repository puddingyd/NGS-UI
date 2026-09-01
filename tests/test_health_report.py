import gzip
import json
import re

from docx import Document

from app import config
from app.adapters import pgx_tsv, snv_tsv
from app.services import docx_export


def test_acmg_ad_narrative_is_risk_not_diagnosis():
    lines = docx_export._health_acmg_narrative(
        "LDLR",
        disease_text="Familial hypercholesterolemia",
        inheritance_codes=["AD"],
        acmg="Pathogenic",
        zygosity="Heterozygous",
        same_gene_count=1,
        sex_karyotype="",
    )

    assert "具有較高的相關疾病風險" in lines[1]
    assert "遺傳諮詢或門診相關專科" in lines[2]
    assert "已罹患" not in "".join(lines)


def test_acmg_ar_single_variant_uses_carrier_wording():
    lines = docx_export._health_acmg_narrative(
        "ATP7B",
        disease_text="Wilson disease",
        inheritance_codes=["AR"],
        acmg="Likely pathogenic",
        zygosity="Heterozygous",
        same_gene_count=1,
        sex_karyotype="",
    )

    assert "本次僅檢出一個符合報告條件之變異" in lines[0]
    assert "符合報告條件的變異" not in lines[0]
    assert "檢測結果符合帶因者狀態" in lines[0]
    assert "帶因者" in lines[1]
    assert "多數不會出現典型疾病症狀" not in lines[1]
    assert "然而" not in lines[1]
    assert "故仍可能" in lines[1]
    assert "資料庫尚未收錄" in lines[1]


def test_structure_label_uses_intron_when_exon_is_placeholder():
    assert docx_export._structure_label({
        "exon": ".",
        "intron": "10/10",
    }) == "intron10"
    assert snv_tsv._clean_vep_rank(".") == ""


def test_acmg_ar_multiple_variants_does_not_assume_phase():
    lines = docx_export._health_acmg_narrative(
        "ATP7B",
        disease_text="Wilson disease",
        inheritance_codes=["AR"],
        acmg="Pathogenic",
        zygosity="Heterozygous",
        same_gene_count=2,
        sex_karyotype="",
    )

    assert "體染色體隱性遺傳" in lines[0]
    assert "此為致病性之變異位點。由於兩變異之相位尚未確認" in lines[1]
    assert "此為致病性之變異位點，由於兩變異之相位尚未確認" not in lines[1]
    assert "由於兩變異之相位尚未確認" in lines[1]
    assert "建議進行家族成員檢測，以釐清" in lines[1]
    assert "兩變異是否位於不同等位基因" in lines[1]
    assert "體染色體隱性疾病之雙等位基因致病型態" in lines[1]
    assert "本結果表示受檢者可能具有較高的相關疾病風險" not in lines[1]
    assert "相位分析" in lines[2]
    assert "遺傳諮詢或門診相關專科" in lines[2]


def test_acmg_x_linked_uses_karyotype():
    xy = docx_export._health_acmg_narrative(
        "ABCD1",
        disease_text="X-linked adrenoleukodystrophy",
        inheritance_codes=["XL"],
        acmg="Pathogenic",
        zygosity="Hemizygous",
        same_gene_count=1,
        sex_karyotype="XY",
    )
    xx = docx_export._health_acmg_narrative(
        "ABCD1",
        disease_text="X-linked adrenoleukodystrophy",
        inheritance_codes=["XL"],
        acmg="Pathogenic",
        zygosity="Heterozygous",
        same_gene_count=1,
        sex_karyotype="XX",
    )

    assert "性聯遺傳" in xy[0] and "可能具有較高" in xy[1]
    assert "依本次檢測所判定" not in xy[1]
    assert xx == [
        "ABCD1 基因與「X-linked adrenoleukodystrophy」相關，其遺傳模式為性聯遺傳。",
        "此為致病性之變異位點。女性帶有單一變異時，臨床表現可能受 X 染色體失活型態"
        "及疾病本身表現範圍影響，可能無症狀或出現不同程度之相關表徵。"
        "實際是否發病、發病年齡及疾病嚴重程度亦可能因個人、家族及環境等因素而異。",
        "建議至遺傳諮詢或門診相關專科，結合個人病史、家族史及適當的臨床檢查進一步評估。"
        "必要時可考慮對具血緣關係的家屬進行此特定位點之驗證檢測。",
    ]


def test_acmg_reference_order_matches_display_grouping():
    ids = ["casq2-a", "ryr2", "casq2-b"]
    variants = {
        "casq2-a": {"gene_symbol": "CASQ2"},
        "ryr2": {"gene_symbol": "RYR2"},
        "casq2-b": {"gene_symbol": "CASQ2"},
    }

    assert docx_export._health_acmg_display_order_ids(ids, variants, {}) == [
        "casq2-a",
        "casq2-b",
        "ryr2",
    ]


def test_health_secondary_default_selection_uses_clinvar_only():
    variants = {
        "clinvar-plp": {"CLNSIG": "Likely_pathogenic"},
        "acmg-only": {"CLNSIG": "Likely_benign", "ACMG_classification": "Likely pathogenic"},
    }

    assert docx_export._health_selected_ids(
        {},
        "acmg_sf",
        ["clinvar-plp", "acmg-only"],
        variants,
    ) == ["clinvar-plp"]
    assert docx_export._health_selected_ids(
        {"secondary_findings": {"acmg_sf": {"selected": ["acmg-only"]}}},
        "acmg_sf",
        ["clinvar-plp", "acmg-only"],
        variants,
    ) == ["clinvar-plp", "acmg-only"]


def test_health_selected_panels_merge_once_in_fixed_order():
    variants = {
        "shared": {"CLNSIG": "Pathogenic"},
        "stroke-only": {"CLNSIG": "Pathogenic"},
        "carrier-only": {"CLNSIG": "Pathogenic"},
    }
    categories = {
        "acmg_sf": ["shared"],
        "stroke": ["stroke-only", "shared"],
        "carrier": ["carrier-only"],
    }

    assert docx_export._health_combined_selected_ids(
        {"acmg_sf", "stroke", "carrier"},
        {},
        categories,
        variants,
    ) == ["shared", "stroke-only", "carrier-only"]


def test_health_overlapping_panel_dismissal_wins_globally():
    variants = {"shared": {"CLNSIG": "Pathogenic"}}
    categories = {
        "acmg_sf": ["shared"],
        "stroke": ["shared"],
        "carrier": [],
    }
    report = {
        "secondary_findings": {
            "acmg_sf": {"selected": ["shared"]},
            "stroke": {"dismissed": ["shared"]},
        },
    }

    assert docx_export._health_combined_selected_ids(
        {"acmg_sf", "stroke"}, report, categories, variants,
    ) == []


def test_health_non_acmg_pure_ar_single_heterozygous_is_carrier():
    variants = {
        "carrier-variant": {
            "gene_symbol": "TESTAR",
            "zygosity": "Heterozygous",
            "Disease1": "Test recessive disease (AR)",
        },
    }

    assert docx_export._health_acmg_categorized_ids(
        ["carrier-variant"], variants, {},
    ) == ([], ["carrier-variant"])


def test_health_karyotype_prefers_ploidy_sidecar(tmp_path, monkeypatch):
    sample_dir = tmp_path / "S1"
    sample_dir.mkdir()
    with gzip.open(sample_dir / "ploidy.vcf.gz", "wt", encoding="utf-8") as handle:
        handle.write("##fileformat=VCFv4.2\n##estimatedSexKaryotype=XY\n#CHROM\tPOS\n")
    monkeypatch.setattr(config, "LEGACY_TERTIARY_OUTPUT_ROOT", tmp_path)

    assert docx_export._health_sex_karyotype("S1", {"Sex": "F"}) == "XY"


def test_health_karyotype_does_not_read_uncopied_source_sibling(tmp_path, monkeypatch):
    sample_dir = tmp_path / "S1"
    source_dir = tmp_path / "source"
    sample_dir.mkdir()
    source_dir.mkdir()
    source_vcf = source_dir / "VAL-3.hard-filtered.vcf.gz"
    source_vcf.touch()
    with gzip.open(source_dir / "VAL-3.ploidy.vcf.gz", "wt", encoding="utf-8") as handle:
        handle.write("##estimatedSexKaryotype=XX\n#CHROM\tPOS\n")
    (sample_dir / "pipeline_source.json").write_text(
        json.dumps({"source_vcf_path": str(source_vcf)}),
        encoding="utf-8",
    )
    monkeypatch.setattr(config, "LEGACY_TERTIARY_OUTPUT_ROOT", tmp_path)

    assert docx_export._health_sex_karyotype("S1", {"Sex": "M"}) == "XY"


def test_health_karyotype_preserves_nonstandard_call(tmp_path, monkeypatch):
    sample_dir = tmp_path / "S1"
    sample_dir.mkdir()
    with gzip.open(sample_dir / "ploidy.vcf.gz", "wt", encoding="utf-8") as handle:
        handle.write("##estimatedSexKaryotype=XXY\n#CHROM\tPOS\n")
    monkeypatch.setattr(config, "LEGACY_TERTIARY_OUTPUT_ROOT", tmp_path)

    assert docx_export._health_sex_karyotype("S1", {"Sex": "M"}) == "XXY"


def test_pgx_summary_prefers_cpic_for_same_gene_and_drug():
    pgx = {
        "guideline_annotations": [
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "alternate_drug_available": True,
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "Use prasugrel or ticagrelor at standard dose.",
            },
            {
                "section": "FDA PGx Association",
                "fda_category": "therapeutic_management",
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "Consider an alternative.",
            },
        ],
    }

    alerts = docx_export._pgx_summary_alerts(pgx)

    assert [row["source"] for row in alerts] == ["CPIC"]
    assert alerts[0]["recommendation"] == "Use prasugrel or ticagrelor at standard dose."


def test_pgx_report_genes_are_fixed_cpic_level_a_scope():
    genes = docx_export._pgx_report_genes({})

    assert len(genes) == 21
    assert "VKORC1" in genes
    assert "NAT2" in genes
    assert "IFNL3" not in genes
    assert "CYP3A4" not in genes


def test_pgx_full_groups_include_cpic_and_fda_for_json_only_gene():
    pgx = {
        "genes": {
            "VKORC1": {
                "additional": True,
                "details": {
                    "label": "rs9923231 variant (T)/rs9923231 variant (T)",
                    "phenotypes": ["-1639 AA"],
                    "allele1_function": "Higher coumarin sensitivity",
                    "allele2_function": "Higher coumarin sensitivity",
                },
            },
        },
        "guideline_annotations": [
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "dosing_information": True,
                "genes": ["VKORC1"],
                "drug": "warfarin",
                "recommendation": "Use genotype-guided dosing.",
            },
            {
                "section": "FDA PGx Association",
                "fda_category": "therapeutic_management",
                "classification": "Unspecified",
                "genes": ["VKORC1"],
                "drug": "warfarin",
                "recommendation": "Select initial dosage using clinical and genetic factors.",
            },
        ],
    }

    assert docx_export._pgx_full_groups(pgx) == [{
        "gene": "VKORC1",
        "diplotype": "rs9923231 T/T（-1639 A/A）",
        "phenotype": "Higher coumarin sensitivity",
        "recommendations": [
            {
                "drug": "warfarin",
                "source": "CPIC",
                "recommendation": "Use genotype-guided dosing.",
                "level": "Strong",
            },
            {
                "drug": "warfarin",
                "source": "FDA PGx Association",
                "recommendation": "Select initial dosage using clinical and genetic factors.",
                "level": "Therapeutic Management",
            },
        ],
    }]


def test_pgx_summary_excludes_standard_cpic_action():
    pgx = {
        "guideline_annotations": [{
            "section": "CPIC Guideline Annotation",
            "classification": "Strong",
            "dosing_information": False,
            "alternate_drug_available": False,
            "other_prescribing_guidance": False,
            "genes": ["DPYD"],
            "drug": "fluorouracil",
            "recommendation": "Use label-recommended dosage and administration.",
        }],
    }

    assert docx_export._pgx_summary_alerts(pgx) == []


def test_pgx_not_recommended_is_actionable_alternative():
    assert docx_export._pgx_action_zh("Ivacaftor is not recommended") == "考慮替代藥物"


def test_pgx_full_groups_exclude_standard_and_uncertain_results():
    pgx = {
        "genes": {
            "DPYD": {"details": {"label": "Reference/Reference", "phenotypes": ["Normal Metabolizer"]}},
            "CACNA1S": {"details": {"label": "Reference/Reference", "phenotypes": ["Uncertain Susceptibility"]}},
        },
        "guideline_annotations": [
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "genes": ["DPYD"],
                "drug": "fluorouracil",
                "recommendation": "Use label-recommended dosage and administration.",
            },
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "other_prescribing_guidance": True,
                "genes": ["CACNA1S"],
                "drug": "sevoflurane",
                "recommendation": "Clinical findings should guide use.",
            },
        ],
    }

    assert docx_export._pgx_full_groups(pgx) == []


def test_pgx_full_groups_deduplicate_same_recommendation_at_strongest_level():
    base = {
        "section": "CPIC Guideline Annotation",
        "alternate_drug_available": True,
        "genes": ["CYP2C19"],
        "drug": "clopidogrel",
        "recommendation": "Avoid clopidogrel if possible.",
    }
    pgx = {
        "genes": {
            "CYP2C19": {"details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]}},
        },
        "guideline_annotations": [
            {**base, "classification": "Moderate"},
            {**base, "classification": "Strong"},
        ],
    }

    groups = docx_export._pgx_full_groups(pgx)

    assert len(groups[0]["recommendations"]) == 1
    assert groups[0]["recommendations"][0]["level"] == "Strong"


def test_pgx_genotype_rows_use_fixed_cpic_level_a_scope():
    pgx = {
        "genes": {
            "CYP2C19": {"details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]}},
            "CYP3A4": {"details": {"label": "*1/*1", "phenotypes": ["Normal Metabolizer"]}},
        },
    }

    rows = docx_export._pgx_genotype_rows(pgx)

    assert len(rows) == 21
    assert ["CYP2C19", "*2/*2", "Poor Metabolizer"] in rows
    assert not any(row[0] == "IFNL3" for row in rows)
    assert not any(row[0] == "CYP3A4" for row in rows)
    assert any(row[0] == "ABCG2" and row[1:] == ["—", "No phenotype assigned"] for row in rows)


def test_pgx_hla_screening_rows_follow_reportable_allele_order():
    pgx = {
        "genes": {
            "HLA-A": {
                "details": {
                    "label": "A*11:01/A*02:06",
                    "phenotypes": ["HLA-A*31:01 negative"],
                },
            },
            "HLA-B": {
                "details": {
                    "label": "B*40:01/B*15:02",
                    "phenotypes": [
                        "HLA-B*15:02 positive",
                        "HLA-B*57:01 negative",
                        "HLA-B*58:01 negative",
                    ],
                },
            },
        },
    }

    assert docx_export._pgx_hla_screening_rows(pgx) == [
        {
            "parent_gene": "HLA-A", "gene": "HLA-A*31:01",
            "allele1": "Negative", "allele2": "Negative",
            "genotype": "Negative/Negative",
        },
        {
            "parent_gene": "HLA-B", "gene": "HLA-B*15:02",
            "allele1": "Negative", "allele2": "Positive",
            "genotype": "Negative/Positive",
        },
        {
            "parent_gene": "HLA-B", "gene": "HLA-B*57:01",
            "allele1": "Negative", "allele2": "Negative",
            "genotype": "Negative/Negative",
        },
        {
            "parent_gene": "HLA-B", "gene": "HLA-B*58:01",
            "allele1": "Negative", "allele2": "Negative",
            "genotype": "Negative/Negative",
        },
    ]


def test_pgx_hla_screening_rows_can_derive_status_from_called_genotype():
    pgx = {
        "genes": {
            "HLA-A": {
                "details": {
                    "label": "",
                    "allele1_name": "A*31:01",
                    "allele2_name": "A*02:06",
                },
            },
            "HLA-B": {
                "details": {
                    "label": "",
                    "allele1_name": "B*40:01",
                    "allele2_name": "B*15:02",
                },
            },
        },
    }

    rows = docx_export._pgx_hla_screening_rows(pgx)

    assert [(row["allele1"], row["allele2"]) for row in rows] == [
        ("Positive", "Negative"),
        ("Negative", "Positive"),
        ("Negative", "Negative"),
        ("Negative", "Negative"),
    ]


def test_pgx_hla_screening_rows_do_not_treat_missing_calls_as_negative():
    rows = docx_export._pgx_hla_screening_rows({"genes": {}})

    assert {
        (row["allele1"], row["allele2"])
        for row in rows
    } == {("No Result", "No Result")}


def test_pgx_hla_aggregate_positive_does_not_invent_two_positive_copies():
    pgx = {
        "genes": {
            "HLA-B": {
                "details": {
                    "label": "*15:02 negative；*57:01 negative；*58:01 positive",
                    "phenotypes": [
                        "*15:02 negative；*57:01 negative；*58:01 positive",
                    ],
                },
            },
        },
    }

    rows = {
        row["gene"]: (row["allele1"], row["allele2"])
        for row in docx_export._pgx_hla_screening_rows(pgx)
    }

    assert rows["HLA-B*15:02"] == ("Negative", "Negative")
    assert rows["HLA-B*58:01"] == ("Positive", "No Result")


def test_pgx_health_genotype_rows_use_consistent_allele_notation():
    pgx = {
        "genes": {
            "ABCG2": {"details": {
                "allele1_name": "rs2231142 reference (G)",
                "allele2_name": "rs2231142 variant (T)",
            }},
            "CYP2D6": {"details": {
                "allele1_name": "*36+*10",
                "allele2_name": "*36+*10",
            }},
            "CFTR": {"details": {
                "label": "Reference/Reference",
                "allele1_name": "ivacaftor non-responsive CFTR sequence",
                "allele2_name": "ivacaftor non-responsive CFTR sequence",
            }},
            "DPYD": {"details": {
                "allele1_name": "c.85T>C(*9A)",
                "allele2_name": "reference",
            }},
            "G6PD": {"details": {"label": "B (reference)"}},
            "HLA-B": {"details": {
                "allele1_name": "B*58:01",
                "allele2_name": "B*58:01",
            }},
            "MT-RNR1": {"details": {"label": "reference"}},
            "VKORC1": {"details": {
                "allele1_name": "rs9923231 reference (C)",
                "allele2_name": "rs9923231 variant (T)",
            }},
        },
    }

    rows = docx_export._pgx_health_genotype_rows(pgx)
    by_test = {row["test"]: row for row in rows}

    assert len(rows) == 25
    assert by_test["ABCG2 rs2231142 (c.421C>A)"] == {
        "test": "ABCG2 rs2231142 (c.421C>A)",
        "allele1": "C (Reference)",
        "allele2": "A (Variant)",
    }
    assert by_test["CYP2D6"] == {
        "test": "CYP2D6", "allele1": "*36+*10", "allele2": "*36+*10",
    }
    assert by_test["CFTR"] == {
        "test": "CFTR", "allele1": "Reference", "allele2": "Reference",
    }
    assert by_test["DPYD"] == {
        "test": "DPYD", "allele1": "c.85T>C (*9A)", "allele2": "Reference",
    }
    assert by_test["G6PD"] == {
        "test": "G6PD", "allele1": "B (Reference)", "allele2": "N/A",
    }
    assert by_test["MT-RNR1"] == {
        "test": "MT-RNR1", "allele1": "Reference", "allele2": "N/A",
    }
    assert by_test["VKORC1 rs9923231 (c.-1639G>A)"] == {
        "test": "VKORC1 rs9923231 (c.-1639G>A)",
        "allele1": "G (Reference)",
        "allele2": "A (Variant)",
    }
    assert by_test["HLA-B"] == {
        "test": "HLA-B", "allele1": "B*58:01", "allele2": "B*58:01",
    }
    assert by_test["  HLA-B*58:01"] == {
        "test": "  HLA-B*58:01", "allele1": "Positive", "allele2": "Positive",
    }
    assert by_test["  HLA-B*15:02"] == {
        "test": "  HLA-B*15:02", "allele1": "Negative", "allele2": "Negative",
    }
    assert [
        row["test"] for row in rows
        if row["test"].strip() in {
            "HLA-A*31:01", "HLA-B*15:02", "HLA-B*57:01", "HLA-B*58:01",
        }
    ] == [
        "  HLA-A*31:01",
        "  HLA-B*15:02",
        "  HLA-B*57:01",
        "  HLA-B*58:01",
    ]
    assert by_test["CACNA1S"] == {
        "test": "CACNA1S", "allele1": "No Result", "allele2": "No Result",
    }


def test_pgx_health_dpyd_multiple_unphased_sources_span_both_allele_columns():
    pgx = {
        "genes": {
            "DPYD": {
                "details": {
                    "effectively_phased": False,
                    "source_diplotypes": [
                        {"allele1_name": "c.1627A>G(*5)", "allele2_name": ""},
                        {"allele1_name": "c.1896T>C", "allele2_name": ""},
                    ],
                    "variants": [
                        {
                            "call": "T/C", "reference": "T",
                            "allele_names": ["c.1627A>G(*5)"],
                            "phased": False, "phase_set": "",
                        },
                        {
                            "call": "A/G", "reference": "A",
                            "allele_names": ["c.1896T>C"],
                            "phased": False, "phase_set": "",
                        },
                    ],
                },
            },
        },
    }

    rows = docx_export._pgx_health_genotype_rows(pgx)
    dpyd = next(row for row in rows if row["test"].startswith("DPYD"))

    assert dpyd == {
        "test": "DPYD（相位未定）",
        "allele1": "",
        "allele2": "",
        "result_span": "檢出變異：c.1627A>G (*5) (het)；c.1896T>C (het)",
    }

    doc = Document()
    docx_export._render_health_pgx_section(doc, "藥物基因體學", pgx)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    normalized = re.sub(r"\s+", " ", text)
    assert "DPYD（相位未定）" in normalized
    assert "檢出變異：c.1627A>G (*5) (het)；c.1896T>C (het)" in normalized
    assert "DPYD（相位未定） No Result" not in normalized


def test_health_pgx_genotype_table_inserts_hla_screens_after_parent_gene():
    doc = Document()
    pgx = {
        "genes": {
            "HLA-A": {
                "details": {
                    "label": "A*11:01/A*02:06",
                    "phenotypes": ["*31:01 negative"],
                },
            },
            "HLA-B": {
                "details": {
                    "label": "B*40:01/B*15:02",
                    "phenotypes": [
                        "*15:02 positive",
                        "*57:01 negative",
                        "*58:01 negative",
                    ],
                },
            },
        },
    }

    docx_export._render_health_pgx_section(doc, "藥物基因體學", pgx)

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    normalized_lines = [re.sub(r"\s+", " ", line.strip()) for line in text.splitlines()]
    assert normalized_lines.index("HLA-A A*11:01 A*02:06") < normalized_lines.index(
        "HLA-A*31:01 Negative Negative"
    )
    assert normalized_lines.index("HLA-A*31:01 Negative Negative") < normalized_lines.index(
        "HLA-B B*40:01 B*15:02"
    )
    assert normalized_lines.index("HLA-B B*40:01 B*15:02") < normalized_lines.index(
        "HLA-B*15:02 Negative Positive"
    )
    assert normalized_lines.index("HLA-B*15:02 Negative Positive") < normalized_lines.index(
        "HLA-B*57:01 Negative Negative"
    ) < normalized_lines.index("HLA-B*58:01 Negative Negative")

    raw_lines = text.splitlines()
    parent_a = next(line for line in raw_lines if line.lstrip().startswith("HLA-A "))
    parent_b = next(line for line in raw_lines if line.lstrip().startswith("HLA-B "))
    for test_name, parent_line in (
        ("HLA-A*31:01", parent_a),
        ("HLA-B*15:02", parent_b),
        ("HLA-B*57:01", parent_b),
        ("HLA-B*58:01", parent_b),
    ):
        child_line = next(
            line for line in raw_lines if line.lstrip().startswith(f"{test_name} ")
        )
        assert len(child_line) - len(child_line.lstrip()) == (
            len(parent_line) - len(parent_line.lstrip()) + 2
        )


def test_health_pgx_excludes_ifnl3_from_all_sections_and_gene_list():
    pgx = {
        "genes": {
            "CYP2C19": {
                "details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]},
            },
            "IFNL3": {
                "details": {
                    "label": "rs12979860 C/C",
                    "phenotypes": ["Favorable response"],
                },
            },
        },
        "guideline_annotations": [
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "alternate_drug_available": True,
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "Use prasugrel or ticagrelor at standard dose.",
            },
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "alternate_drug_available": True,
                "genes": ["IFNL3"],
                "drug": "peginterferon alfa-2a",
                "recommendation": "Consider an alternative treatment.",
            },
        ],
    }
    doc = Document()

    groups = docx_export._render_health_pgx_section(doc, "藥物基因體學", pgx)
    docx_export._section_health_annotations(doc, {"pgx"}, pgx)
    docx_export._render_health_pgx_appendix(doc, "附錄一、完整用藥建議", groups)

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "基因型" in text
    assert "完整用藥建議" in text
    assert "用藥建議概覽" not in text
    assert "藥物建議摘要" not in text
    assert "基因型與表現型" not in text
    genotype_table = next(
        paragraph.text for paragraph in doc.paragraphs
        if "基因／檢測項目" in paragraph.text and "Allele 1" in paragraph.text
    )
    assert "表型" not in genotype_table
    assert "CYP2C19" in text and "Clopidogrel" in text
    assert "IFNL3" not in text
    assert "peginterferon" not in text.lower()


def test_pgx_full_recommendation_table_width_matches_genotype_table():
    columns = docx_export._PGX_FULL_RECOMMENDATION_COLUMNS

    assert columns == [
        ("藥物", 23, "word-buffered"),
        ("基因與表型", 18, "word-buffered"),
        ("CPIC/FDA 建議", 43, "word-buffered"),
    ]
    assert sum(col[1] for col in columns) + 2 == 86
    assert docx_export._wrap_to_cols(
        "Amifampridine phosphate",
        24,
        mode="word-buffered",
    ) == ["Amifampridine phosphate"]
    assert docx_export._wrap_to_cols(
        "FDA Therapeutic Management",
        27,
        mode="buffered",
    ) == ["FDA Therapeutic Management"]


def test_pgx_multigene_recommendation_only_lists_actionable_gene():
    pgx = {
        "genes": {
            "NUDT15": {
                "details": {"label": "*1/*3", "phenotypes": ["Intermediate Metabolizer"]},
            },
            "TPMT": {
                "details": {"label": "*1/*1", "phenotypes": ["Normal Metabolizer"]},
            },
        },
        "guideline_annotations": [{
            "section": "CPIC Guideline Annotation",
            "classification": "Moderate",
            "dosing_information": True,
            "genes": ["NUDT15", "TPMT"],
            "drug": "thioguanine",
            "recommendation": "Reduce the starting dose and monitor blood counts.",
        }],
    }

    groups = docx_export._pgx_drug_groups(pgx)
    rows = docx_export._pgx_summary_rows_from_drug_groups(groups)

    assert list(groups[0]["genes"]) == ["NUDT15"]
    assert {rec["gene"] for rec in groups[0]["recommendations"]} == {"NUDT15"}
    assert rows[0]["gene"] == "NUDT15"


def test_pgx_gene_result_falls_back_to_tsv_for_mt_rnr1():
    pgx = {
        "genes": {
            "MT-RNR1": {
                "diplotype": "m.1555A>G positive",
                "phenotype": "HIGH risk",
                "details": {
                    "label": "Unknown",
                    "phenotypes": ["No Result"],
                },
            },
        },
    }

    assert docx_export._pgx_gene_result(pgx, "MT-RNR1") == ("m.1555A>G positive", "HIGH risk")


def test_pgx_gene_result_does_not_fall_back_to_tsv_outside_mt_rnr1():
    pgx = {
        "genes": {
            "CYP2C19": {
                "diplotype": "*2/*2",
                "phenotype": "Poor Metabolizer",
                "details": {},
            },
        },
    }

    assert docx_export._pgx_gene_result(pgx, "CYP2C19") == ("", "")


def test_pgx_report_view_is_shared_docx_projection():
    pgx = {
        "genes": {
            "CYP2C19": {
                "details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]},
            },
        },
        "guideline_annotations": [{
            "section": "CPIC Guideline Annotation",
            "classification": "Strong",
            "alternate_drug_available": True,
            "genes": ["CYP2C19"],
            "drug": "clopidogrel",
            "recommendation": "Use an alternative antiplatelet agent.",
        }],
    }

    view = docx_export.build_pgx_report_view(pgx)

    assert len(view["report_genes"]) == 21
    assert view["summary_rows"] == [{
        "drug": "Clopidogrel",
        "gene": "CYP2C19",
        "action": "考慮替代藥物",
        "source_level": "CPIC Strong",
    }]
    assert view["action_categories"] == [{
        "action": "考慮替代藥物",
        "drugs": ["Clopidogrel"],
    }]
    assert view["full_recommendations"][0]["gene_phenotype"] == "CYP2C19 Poor Metabolizer"
    assert view["summary"]["actionable_genes"] == ["CYP2C19"]
    cyp2c19 = next(row for row in view["analysis_genes"] if row["gene"] == "CYP2C19")
    assert cyp2c19["requires_attention"] is True
    assert cyp2c19["related_drugs"] == ["Clopidogrel"]
    assert cyp2c19["recommendation_rows"][0]["recommendation"].endswith("(CPIC Strong)")
    assert [row["action"] for row in view["analysis_action_categories"]] == [
        "調整劑量並監測",
        "考慮替代藥物",
        "加強不良反應監測",
        "使用前確認表型或檢驗",
        "其他用藥建議（參考完整建議與最新藥品仿單）",
        "不需調整／無明確調整建議",
    ]
    assert view["analysis_action_categories"][0]["drugs"] == []


def test_pgx_analysis_lists_related_drugs_for_normal_genes():
    pgx = {
        "genes": {
            "DPYD": {
                "details": {"label": "*1/*1", "phenotypes": ["Normal Metabolizer"]},
            },
        },
        "guideline_annotations": [{
            "section": "CPIC Guideline Annotation",
            "classification": "Strong",
            "genes": ["DPYD"],
            "drug": "fluorouracil",
            "recommendation": "Use standard dosing.",
        }],
    }

    view = docx_export.build_pgx_report_view(pgx)

    dpyd = next(row for row in view["analysis_genes"] if row["gene"] == "DPYD")
    assert dpyd["requires_attention"] is False
    assert dpyd["related_drugs"] == ["Fluorouracil"]
    assert dpyd["recommendation_rows"] == [{
        "drug": "Fluorouracil",
        "recommendation": "目前基因型／表型未符合本報告需調整用藥的條件。",
        "requires_adjustment": False,
    }]
    assert view["analysis_action_categories"][-1]["drugs"] == ["Fluorouracil"]


def test_pgx_analysis_orders_adjustment_before_no_adjustment():
    pgx = {
        "genes": {
            "CYP2C19": {
                "details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]},
            },
        },
        "guideline_annotations": [
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "alternate_drug_available": True,
                "genes": ["CYP2C19"],
                "drug": "z-drug",
                "recommendation": "Use an alternative treatment.",
            },
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Optional",
                "genes": ["CYP2C19"],
                "drug": "a-drug",
                "recommendation": "Consider the clinical context.",
            },
        ],
    }

    view = docx_export.build_pgx_report_view(pgx)

    cyp2c19 = next(row for row in view["analysis_genes"] if row["gene"] == "CYP2C19")
    assert [row["drug"] for row in cyp2c19["recommendation_rows"]] == ["Z-drug", "A-drug"]
    assert [row["requires_adjustment"] for row in cyp2c19["recommendation_rows"]] == [True, False]


def test_pgx_overview_maps_generic_fallback_to_clear_other_category():
    groups = [{
        "drug": "Example drug",
        "action": "參考最新藥品仿單",
        "recommendations": [{
            "gene": "CYP2C19",
            "source": "CPIC",
            "level": "Moderate",
            "action": "參考最新藥品仿單",
            "source_priority": 1,
        }],
    }]

    assert docx_export._pgx_action_categories(groups) == [(
        "其他用藥建議（參考完整建議與最新藥品仿單）",
        ["Example drug"],
    )]
    assert docx_export._pgx_summary_rows_from_drug_groups(groups)[0]["action"] == (
        "其他用藥建議（參考完整建議與最新藥品仿單）"
    )
    assert docx_export._pgx_action_zh("Confirm phenotype before treatment.") == "使用前確認表型或檢驗"


def test_pgx_analysis_overview_covers_each_json_drug_once():
    pgx = {
        "genes": {
            "CYP2C19": {
                "details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]},
            },
        },
        "guideline_annotations": [
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "alternate_drug_available": True,
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "Use an alternative antiplatelet agent.",
            },
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Moderate",
                "other_prescribing_guidance": True,
                "genes": ["CYP2C19"],
                "drug": "fallback-drug",
                "recommendation": "Consider the patient's overall clinical context.",
            },
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Optional",
                "genes": ["CYP2C19"],
                "drug": "routine-drug",
                "recommendation": "Use usual care.",
            },
            {
                "section": "FDA Label Annotation",
                "classification": "Unspecified",
                "other_prescribing_guidance": True,
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "See the current prescribing information.",
            },
        ],
    }

    view = docx_export.build_pgx_report_view(pgx)
    by_action = {
        row["action"]: row["drugs"]
        for row in view["analysis_action_categories"]
    }

    assert by_action["考慮替代藥物"] == ["Clopidogrel"]
    assert by_action["其他用藥建議（參考完整建議與最新藥品仿單）"] == ["Fallback-drug"]
    assert by_action["不需調整／無明確調整建議"] == ["Routine-drug"]
    categorized = [
        drug
        for category in view["analysis_action_categories"]
        for drug in category["drugs"]
    ]
    assert sorted(categorized, key=str.casefold) == [
        "Clopidogrel",
        "Fallback-drug",
        "Routine-drug",
    ]


def test_pgx_analysis_overview_allows_only_mt_rnr1_tsv_drug_supplement():
    pgx = {
        "genes": {
            "MT-RNR1": {
                "diplotype": "m.1555A>G positive",
                "phenotype": "HIGH risk",
                "drugs": [{"drug": "aminoglycosides", "recommendations": []}],
                "details": {"label": "Unknown", "phenotypes": ["No Result"]},
            },
            "CYP2C19": {
                "diplotype": "*2/*2",
                "phenotype": "Poor Metabolizer",
                "drugs": [{"drug": "legacy-tsv-drug", "recommendations": []}],
                "details": {},
            },
        },
    }

    view = docx_export.build_pgx_report_view(pgx)

    assert view["analysis_action_categories"][-1] == {
        "action": "不需調整／無明確調整建議",
        "drugs": ["Aminoglycosides"],
    }


def test_pgx_genotype_rows_treat_dash_phenotype_as_not_assigned():
    pgx = {
        "genes": {
            "VKORC1": {
                "details": {
                    "label": "rs9923231 T/T",
                    "phenotypes": ["—"],
                },
            },
        },
    }

    rows = docx_export._pgx_genotype_rows(pgx)

    assert ["VKORC1", "rs9923231 T/T", "No phenotype assigned"] in rows


def test_pgx_gene_result_falls_back_to_allele_function_for_cyp4f2():
    pgx = {
        "genes": {
            "CYP4F2": {
                "details": {
                    "label": "*1/*3",
                    "phenotypes": ["n/a"],
                    "allele1_function": "Normal function",
                    "allele2_function": "Decreased function",
                },
            },
        },
    }

    assert docx_export._pgx_gene_result(pgx, "CYP4F2") == (
        "*1/*3",
        "Normal function；Decreased function",
    )


def test_pgx_gene_result_keeps_vkorc1_sensitivity_as_phenotype():
    pgx = {
        "genes": {
            "VKORC1": {
                "details": {
                    "label": "rs9923231 variant (T)/rs9923231 variant (T)",
                    "phenotypes": ["-1639 AA"],
                    "allele1_function": "Higher coumarin sensitivity",
                    "allele2_function": "Higher coumarin sensitivity",
                },
            },
        },
    }

    assert docx_export._pgx_gene_result(pgx, "VKORC1") == (
        "rs9923231 T/T（-1639 A/A）",
        "Higher coumarin sensitivity",
    )


def test_pgx_drug_groups_are_drug_centric_and_keep_strongest_cpic():
    pgx = {
        "genes": {
            "CYP2C19": {"details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]}},
        },
        "guideline_annotations": [
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Moderate",
                "alternate_drug_available": True,
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "Avoid clopidogrel if possible. Consider an alternative P2Y12 inhibitor.",
            },
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "alternate_drug_available": True,
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "Avoid clopidogrel if possible. Use prasugrel or ticagrelor.",
            },
            {
                "section": "FDA Label Annotation",
                "alternate_drug_available": True,
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "Consider use of another platelet P2Y12 inhibitor.",
            },
            {
                "section": "FDA PGx Association",
                "fda_category": "therapeutic_management",
                "genes": ["CYP2C19"],
                "drug": "clopidogrel",
                "recommendation": "Consider use of another platelet P2Y12 inhibitor.",
            },
        ],
    }

    groups = docx_export._pgx_drug_groups(pgx)

    assert len(groups) == 1
    assert groups[0]["drug"] == "Clopidogrel"
    assert groups[0]["action"] == "考慮替代藥物"
    assert groups[0]["source_level"] == "CPIC Strong"
    assert docx_export._pgx_gene_phenotype_text(groups[0]) == "CYP2C19 Poor Metabolizer"
    assert [
        (rec["source"], rec["level"])
        for rec in groups[0]["recommendations"]
    ] == [
        ("CPIC", "Strong"),
        ("FDA PGx Association", "Therapeutic Management"),
        ("FDA Label", "Unspecified"),
    ]


def test_pgx_recommendation_text_combines_fda_sources_after_content():
    text = docx_export._pgx_recommendation_text({
        "recommendations": [
            {
                "source": "CPIC",
                "level": "Strong",
                "recommendation": '"Use alternative therapy."',
            },
            {
                "source": "FDA Label",
                "level": "Unspecified",
                "recommendation": '"Use label dosing."',
            },
            {
                "source": "FDA PGx Association",
                "level": "Therapeutic Management",
                "recommendation": '"Adjust dose."',
            },
        ],
    })

    assert text == (
        "Use alternative therapy. (CPIC Strong)；"
        "Adjust dose. Use label dosing. (FDA Therapeutic Management / FDA Label)"
    )
    assert '"' not in text

    cleaned = docx_export._pgx_clean_recommendation_text(
        '&quot;[...]increase monitoring for adverse reactions.&quot; '
        'In CYP2C19 poor metabolizers ... the starting dose should be reduced. '
        'No dosage adjustment is needed...For known poor metabolizers, '
        '...thioridazine is contraindicated...in patients.'
    )
    assert cleaned == (
        "increase monitoring for adverse reactions. "
        "In CYP2C19 poor metabolizers the starting dose should be reduced. "
        "No dosage adjustment is needed. For known poor metabolizers, "
        "thioridazine is contraindicated in patients."
    )
    assert "..." not in cleaned and "[...]" not in cleaned


def test_pgx_summary_rows_exclude_fda_label_only_drugs():
    rows = docx_export._pgx_summary_rows_from_drug_groups([
        {
            "drug": "Belzutifan",
            "genes": {"CYP2C19": {}},
            "action": "調整劑量並監測",
            "source_level": "FDA Label",
            "recommendations": [
                {
                    "gene": "CYP2C19",
                    "source": "FDA Label",
                    "level": "Unspecified",
                    "action": "調整劑量並監測",
                    "source_priority": 3,
                },
            ],
        },
        {
            "drug": "Deutetrabenazine",
            "genes": {"CYP2D6": {}},
            "action": "調整劑量並監測",
            "source_level": "FDA Therapeutic Management",
            "recommendations": [
                {
                    "gene": "CYP2D6",
                    "source": "FDA PGx Association",
                    "level": "Therapeutic Management",
                    "action": "調整劑量並監測",
                    "source_priority": 2,
                },
            ],
        },
    ])

    assert rows == [{
        "drug": "Deutetrabenazine",
        "gene": "CYP2D6",
        "action": "調整劑量並監測",
        "source_level": "FDA Therapeutic Management",
    }]


def test_pgx_summary_reuses_overview_primary_action_and_source():
    rows = docx_export._pgx_summary_rows_from_drug_groups([{
        "drug": "Eliglustat",
        "genes": {"CYP2D6": {}},
        "action": "調整劑量並監測",
        "source_level": "FDA Label",
        "recommendations": [
            {
                "gene": "CYP2D6",
                "source": "FDA PGx Association",
                "level": "Therapeutic Management",
                "action": "考慮替代藥物",
                "source_priority": 2,
            },
            {
                "gene": "CYP2D6",
                "source": "FDA Label",
                "level": "Unspecified",
                "action": "調整劑量並監測",
                "source_priority": 3,
            },
        ],
    }])

    assert rows == [{
        "drug": "Eliglustat",
        "gene": "CYP2D6",
        "action": "調整劑量並監測",
        "source_level": "FDA Label",
    }]


def test_pgx_adapter_keeps_json_when_tsv_is_absent(tmp_path, monkeypatch):
    monkeypatch.setattr(pgx_tsv, "_compact_report_json", lambda _path: {
        "pharmcat_version": "3.2.0",
        "data_version": "2026-02-09",
        "timestamp": "2026-06-30",
        "messages": [],
        "genes": {
            "VKORC1": {
                "label": "rs9923231 variant (T)/rs9923231 variant (T)",
                "phenotypes": ["-1639 AA"],
            },
        },
        "guideline_annotations": [{
            "section": "FDA PGx Association",
            "genes": ["VKORC1"],
        }],
    })

    result = pgx_tsv.load_pgx(
        tmp_path / "missing.pgx.tsv",
        tmp_path / "pharmcat.report.json",
    )

    assert result["pharmcat_version"] == "3.2.0"
    assert result["additional_genes"] == ["VKORC1"]
    assert result["genes"]["VKORC1"]["details"]["phenotypes"] == ["-1639 AA"]


def test_pgx_only_annotation_lists_all_cpic_level_a_genes():
    doc = Document()

    docx_export._section_health_annotations(doc, {"pgx"}, {})

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "CPIC) Level A gene" in text
    assert "ABCG2" in text and "VKORC1" in text
    assert "CYP3A4" not in text
    assert "FDA" not in text


def test_health_pathogenicities_translate_and_combine():
    assert docx_export._health_pathogenicities_zh([
        "Pathogenic",
        "Uncertain significance",
    ]) == "致病性及不確定意義"
    assert [
        docx_export._health_pathogenicity_zh(label)
        for label in (
            "Benign",
            "Likely benign",
            "Uncertain significance",
            "Likely pathogenic",
            "Pathogenic",
        )
    ] == ["良性", "疑似良性", "不確定意義", "疑似致病性", "致病性"]


def test_health_acmg_table_keeps_all_five_classes_in_english():
    doc = Document()
    labels = (
        "Benign",
        "Likely benign",
        "Uncertain significance",
        "Likely pathogenic",
        "Pathogenic",
    )
    rows = [
        ({
            "gene_symbol": "TEST",
            "ACMG_classification": label,
            "zygosity": "Heterozygous",
        }, {})
        for label in labels
    ]

    docx_export._health_snv_gene_block(
        doc,
        rows,
        disease_text="Test disease",
        inheritance_codes=["AD"],
        sex_karyotype="",
    )

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    normalized_text = re.sub(r"\s+", " ", text)
    for source in labels:
        assert source in normalized_text
    for translated in ("良性", "疑似良性", "不確定意義", "疑似致病性", "致病性"):
        assert translated in text


def test_health_variant_reference_appendix_translates_acmg_class():
    doc = Document()
    variant = {
        "id": "test-variant",
        "gene_symbol": "TEST",
        "HGVS_C": "c.1A>G",
        "ACMG_classification": "Likely benign",
    }

    docx_export._render_health_variant_reference_appendix(
        doc,
        "附錄一、變異位點參考資料",
        [variant],
        {"edits": {}},
    )

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "評測此變異位點為「疑似良性」" in text
    assert "評測此變異位點為「Likely benign」" not in text


def test_vkorc1_display_moves_position_phenotype_into_genotype():
    assert docx_export._pgx_display_genotype(
        "VKORC1",
        "rs9923231 variant (T)/rs9923231 variant (T)",
        "-1639 AA",
    ) == ("rs9923231 T/T（-1639 A/A）", "—")


def test_pgx_summary_rows_translate_and_sort_by_drug():
    alerts = [
        {
            "gene": "G6PD",
            "drug": "rasburicase",
            "source": "CPIC",
            "level": "Strong",
            "recommendation": "Avoid rasburicase.",
            "priority": True,
        },
        {
            "gene": "G6PD",
            "drug": "dapsone",
            "source": "CPIC",
            "level": "Strong",
            "recommendation": "Avoid dapsone.",
            "priority": True,
        },
    ]

    assert docx_export._pgx_summary_rows(alerts) == [
        {
            "drug": "Dapsone",
            "gene": "G6PD",
            "action": "考慮替代藥物",
            "source_level": "CPIC Strong",
        },
        {
            "drug": "Rasburicase",
            "gene": "G6PD",
            "action": "考慮替代藥物",
            "source_level": "CPIC Strong",
        },
    ]


def test_health_bundle_name_follows_selected_sections():
    assert docx_export._health_test_bundle_name({"acmg_sf", "pgx"}) == (
        "ACMG疾病風險基因及藥物基因體學基因篩檢"
    )
    assert docx_export._health_test_bundle_name({"pgx"}) == "藥物基因體學基因篩檢"
    assert docx_export._health_test_bundle_name({"acmg_sf"}) == "ACMG疾病風險基因篩檢"
    assert docx_export._health_test_bundle_name({"stroke", "carrier"}) == (
        "中風相關基因及帶因者基因篩檢"
    )


def test_health_acmg_section_titles_use_current_wording():
    assert dict(docx_export._HEALTH_SECTION_ORDER)["acmg_sf"] == (
        "第一類：與疾病風險相關之致病性或疑似致病性變異位點"
    )
    assert docx_export._HEALTH_ACMG_GENE_LIST_TITLE.startswith("ACMG疾病風險基因")
    assert "第一類：" not in docx_export._HEALTH_ACMG_GENE_LIST_TITLE
    assert docx_export._NO_HEALTH_CARRIER_VARIANT_TEXT == (
        "於本次檢測之基因中，未檢出疾病資料庫中已收錄且符合帶因者狀態之致病性或疑似致病性變異。"
    )
    assert "可採取醫療處置之疾病風險基因" not in dict(docx_export._HEALTH_SECTION_ORDER)["acmg_sf"]


def test_health_header_and_acmg_caution_follow_current_template():
    doc = Document()

    docx_export._section_health_patient_header(doc, {"LIS_ID": "SF1-dragen"})

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "<<基因醫學部基因檢測檢驗分析研究報告>>" not in text
    assert "<<基因醫學部基因檢測分析研究報告>>" in text
    assert "國立成功大學醫學院附設醫院" in text
    assert "研究報告）" not in text
    assert "次發現基因清單第 3.3 版（2025 年發表）" in docx_export._HEALTH_ACMG_CAUTION
    assert "次發現基因清單 ACMG SF 3.3" not in docx_export._HEALTH_ACMG_CAUTION
    assert "分析與 ACMG SF 3.3 所列遺傳性疾病相關的風險基因" in docx_export._HEALTH_ACMG_CAUTION
    assert "可採取醫療處置之遺傳性疾病" not in docx_export._HEALTH_ACMG_CAUTION


def test_health_pgx_main_body_matches_genotype_only_template():
    doc = Document()
    pgx = {
        "genes": {
            "CYP2C19": {"details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]}},
        },
        "guideline_annotations": [{
            "section": "CPIC Guideline Annotation",
            "classification": "Strong",
            "alternate_drug_available": True,
            "genes": ["CYP2C19"],
            "drug": "clopidogrel",
            "recommendation": "Use an alternative antiplatelet therapy.",
        }],
    }

    groups = docx_export._render_health_pgx_section(doc, "藥物基因體學", pgx)
    docx_export._section_methods(doc, "WGS", health=True)

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]
    intro_index = paragraph_texts.index(
        "  此處列出基因型，完整藥物建議詳見報告末端附錄。若目前使用或未來考慮使用相關藥物，"
        "建議由處方醫師參考下方結果或最新 FDA/CPIC 指引進行評估，"
        "或參考下述官方用藥說明處之連結。"
    )
    assert "其餘未列於下方之藥物" not in text
    assert "其餘未列之藥物" not in text
    assert "本次檢測發現" not in text
    assert "用藥建議概覽" not in text
    assert "藥物建議摘要" not in text
    assert "基因型與表現型" not in text
    assert "參考下方結果或最新 FDA/CPIC 指引進行評估" in text
    assert "或參考下述官方用藥說明處之連結" in text
    assert paragraph_texts[intro_index + 1] == ""
    assert paragraph_texts[intro_index + 2] == "  基因型"
    assert text.index("官方用藥資訊查詢") < text.index("四、檢測方法說明")
    assert "完整用藥建議" not in text
    assert groups


def test_health_pgx_intro_omits_dynamic_drug_summary_and_phenotype():
    doc = Document()
    pgx = {
        "genes": {
            "CYP2C19": {"details": {"label": "*2/*2", "phenotypes": ["Poor Metabolizer"]}},
        },
        "guideline_annotations": [
            {
                "section": "CPIC Guideline Annotation",
                "classification": "Strong",
                "alternate_drug_available": True,
                "genes": ["CYP2C19"],
                "drug": f"drug-{index:02d}",
                "recommendation": "Use an alternative therapy.",
            }
            for index in range(1, 14)
        ] + [{
            "section": "FDA Label Annotation",
            "classification": "Informative PGx",
            "alternate_drug_available": True,
            "genes": ["CYP2C19"],
            "drug": "appendix-only-drug",
            "recommendation": "See the FDA label for prescribing information.",
        }],
    }

    docx_export._render_health_pgx_section(doc, "藥物基因體學", pgx)

    intro = next(
        paragraph.text for paragraph in doc.paragraphs
        if paragraph.text.startswith("  此處列出基因型")
    )
    assert "本次檢測發現" not in intro
    assert "下方摘要表" not in intro
    assert "Drug-13" not in intro
    assert "Appendix-only-drug" not in intro
    assert "表現型" not in intro


def test_health_pgx_hla_positive_wins_over_other_negative_alleles():
    pgx = {
        "genes": {
            "HLA-B": {
                "details": {
                    "label": "*15:02 negative；*57:01 negative；*58:01 positive",
                    "phenotypes": ["*15:02 negative；*57:01 negative；*58:01 positive"],
                },
            },
        },
        "guideline_annotations": [{
            "section": "CPIC Guideline Annotation",
            "classification": "Strong",
            "alternate_drug_available": True,
            "genes": ["HLA-B"],
            "drug": "allopurinol",
            "recommendation": "Avoid allopurinol and use an alternative therapy.",
        }],
    }

    assert docx_export._pgx_gene_has_actionable_result(pgx, "HLA-B") is True
    groups = docx_export._pgx_drug_groups(pgx)
    assert [group["drug"] for group in groups] == ["Allopurinol"]
    assert docx_export._pgx_summary_rows_from_drug_groups(groups) == [{
        "drug": "Allopurinol",
        "gene": "HLA-B",
        "action": "考慮替代藥物",
        "source_level": "CPIC Strong",
    }]


def test_health_acmg_main_body_lists_findings_without_subgroup_catalog():
    doc = Document()
    variants = {
        "ldlr": {
            "id": "ldlr",
            "gene_symbol": "LDLR",
            "ACMG_classification": "Pathogenic",
            "CLNSIG": "Pathogenic",
            "zygosity": "Heterozygous",
        },
    }

    docx_export._render_health_acmg_section(
        doc,
        dict(docx_export._HEALTH_SECTION_ORDER)["acmg_sf"],
        ["ldlr"],
        variants,
        {"edits": {}},
    )

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "LDLR" in text
    assert "血脂相關基因，包含 APOB, LDLR, PCSK9" not in text
    assert text.count(docx_export._NO_HEALTH_VARIANT_TEXT) == 0
    assert text.index(docx_export._HEALTH_ACMG_CAUTION) < text.index(
        "第一類：與疾病風險相關之致病性或疑似致病性變異位點"
    )


def test_health_acmg_categories_split_by_inheritance_count_and_zygosity():
    variants = {
        "ldlr": {"gene_symbol": "LDLR", "zygosity": "Heterozygous"},
        "gla": {"gene_symbol": "GLA", "zygosity": "Heterozygous"},
        "casq2-a": {"gene_symbol": "CASQ2", "zygosity": "Heterozygous"},
        "casq2-b": {"gene_symbol": "CASQ2", "zygosity": "Heterozygous"},
        "btd-hom": {"gene_symbol": "BTD", "zygosity": "Homozygous"},
        "mutyh-carrier": {"gene_symbol": "MUTYH", "zygosity": "Heterozygous"},
    }

    risk_ids, carrier_ids = docx_export._health_acmg_categorized_ids(
        list(variants),
        variants,
        {},
    )

    assert risk_ids == ["ldlr", "casq2-a", "casq2-b", "gla", "btd-hom"]
    assert carrier_ids == ["mutyh-carrier"]


def test_health_acmg_first_class_keeps_requested_negative_when_only_carrier_found():
    doc = Document()
    variants = {
        "mutyh-carrier": {
            "id": "mutyh-carrier",
            "gene_symbol": "MUTYH",
            "ACMG_classification": "Likely pathogenic",
            "CLNSIG": "Likely pathogenic",
            "zygosity": "Heterozygous",
        },
    }

    docx_export._render_health_acmg_section(
        doc,
        dict(docx_export._HEALTH_SECTION_ORDER)["acmg_sf"],
        ["mutyh-carrier"],
        variants,
        {"edits": {}},
    )

    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]
    text = "\n".join(paragraph_texts)
    assert docx_export._NO_HEALTH_VARIANT_TEXT in text
    assert "第二類：符合帶因者狀態之致病性或疑似致病性變異位點" in text
    assert "MUTYH" in text
    assert docx_export._NO_HEALTH_CARRIER_VARIANT_TEXT not in text
    first_title_index = paragraph_texts.index(
        "第一類：與疾病風險相關之致病性或疑似致病性變異位點"
    )
    first_result_index = paragraph_texts.index(f"  {docx_export._NO_HEALTH_VARIANT_TEXT}")
    second_title_index = paragraph_texts.index(
        "第二類：符合帶因者狀態之致病性或疑似致病性變異位點"
    )
    assert first_result_index == first_title_index + 1
    assert paragraph_texts[second_title_index - 1] == ""


def test_health_acmg_both_categories_use_exact_negative_text():
    doc = Document()

    docx_export._render_health_acmg_section(
        doc,
        dict(docx_export._HEALTH_SECTION_ORDER)["acmg_sf"],
        [],
        {},
        {"edits": {}},
    )

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert docx_export._NO_HEALTH_VARIANT_TEXT in text
    assert docx_export._NO_HEALTH_CARRIER_VARIANT_TEXT in text


def test_health_wgs_methods_include_scope_specific_cnv_notes():
    doc = Document()

    docx_export._section_methods(doc, "WGS", health=True)

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "3. 本次檢測平均定序深度 ≧ 27X" in text
    assert "無法檢測出拷貝數變異 (copy number variant)、轉位" in text
    assert "5. 本檢驗採用短讀長全基因體定序" in text
    assert "6. 本檢驗方法並非長片段定序" in text
    assert "相位 (phasing) 及其單套體資訊 (haplotype)" in text
    assert "7. 藥物基因體學分析中，CYP2D6 基因型判定會納入該基因之拷貝數變異" in text
    assert "僅用於 CYP2D6 藥物基因體學判讀" in text
    assert "ACMG SF 或其他基因" not in text
    assert "不代表本檢測已涵蓋其他基因之拷貝數變異" in text
    assert "8. 本實驗方法以次世代方法定序粒線體DNA基因序列" in text
    assert "9. 本檢測報告僅供醫療專業人員參考" in text


def test_health_wes_methods_do_not_claim_wgs_specific_limitations():
    doc = Document()

    docx_export._section_methods(doc, "WES", health=True)

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "短讀長全基因體定序" not in text
    assert "相位 (phasing)" not in text
    assert "5. 藥物基因體學分析中，CYP2D6 基因型判定" in text
    assert "6. 本檢測報告僅供醫療專業人員參考" in text


def test_diagnosis_wgs_methods_keep_existing_depth_and_numbering():
    doc = Document()

    docx_export._section_methods(doc, "WGS", health=False)

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "3. 本次檢測平均定序深度 ≧ 30X" in text
    assert "CYP2D6 基因型判定" not in text
    assert "5. 本實驗方法以次世代方法定序粒線體DNA基因序列" in text
    assert "6. 本檢測報告僅供醫療專業人員參考" in text


def test_health_appendix_orders_acmg_references_before_pgx_recommendations():
    doc = Document()
    variant = {
        "id": "test-variant",
        "gene_symbol": "LDLR",
        "HGVS_C": "c.1A>G",
        "ACMG_classification": "Pathogenic",
    }
    groups = [{
        "drug": "Clopidogrel",
        "genes": {"CYP2C19": {"phenotype": "Poor Metabolizer"}},
        "recommendations": [{
            "source": "CPIC",
            "level": "Strong",
            "recommendation": "Use an alternative antiplatelet agent.",
        }],
    }]

    docx_export._start_health_appendix(doc)
    docx_export._render_health_variant_reference_appendix(
        doc,
        "變異位點參考資料",
        [variant],
        {"edits": {}},
    )
    docx_export._blank(doc)
    docx_export._render_health_pgx_appendix(doc, "完整用藥建議", groups)

    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]
    text = "\n".join(paragraph_texts)
    assert text.index("附錄") < text.index("變異位點參考資料")
    assert text.index("變異位點參考資料") < text.index("完整用藥建議")
    complete_index = paragraph_texts.index("完整用藥建議")
    assert paragraph_texts[complete_index - 1] == ""
    appendix_index = paragraph_texts.index("附錄")
    assert paragraph_texts[appendix_index + 1] == ""
    assert doc.paragraphs[appendix_index].alignment == 1
    assert 'w:type="page"' in doc.element.xml
    assert "附錄一" not in text and "附錄二" not in text


def test_health_gene_list_contains_acmg_subgroups(monkeypatch):
    monkeypatch.setattr(
        docx_export,
        "_health_panel_gene_sections",
        lambda requested: [(docx_export._HEALTH_ACMG_GENE_LIST_TITLE, ["LDLR"])],
    )
    doc = Document()

    docx_export._section_health_annotations(doc, {"acmg_sf"})

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "第一類：ACMG疾病風險基因" not in text
    assert "1. 血脂相關基因，包含 APOB, LDLR, PCSK9" in text
    assert "2. 腫瘤相關基因，包含 APC, BMPR1A" in text
    assert "6. 其它基因，包含 RPE65, TTR" in text


def test_health_pgx_section_includes_visible_official_urls():
    doc = Document()
    docx_export._render_health_pgx_resources(doc)

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    visible_xml_text = "".join(
        node.text or ""
        for node in doc.element.iter()
        if node.tag.endswith("}t")
    )
    targets = {rel.target_ref for rel in doc.part.rels.values() if rel.is_external}
    assert "官方用藥資訊查詢" in text
    assert "以上簡要判讀說明" in text
    assert "未提供完整的劑量調整指示" in text
    assert "可供臨床處方決策參考的藥物基因體資訊" in text
    assert "如下之 CPIC 最新臨床指引" in text
    assert "請務必諮詢臨床醫師或臨床藥理專業人員" in text
    assert docx_export._HEALTH_PGX_RESOURCES[0][0] == "CPIC 最新臨床指引"
    assert {url for _, url in docx_export._HEALTH_PGX_RESOURCES}.issubset(targets)
    for _, url in docx_export._HEALTH_PGX_RESOURCES:
        assert url in visible_xml_text


def test_health_ascii_tables_do_not_emit_keep_next_markers():
    doc = Document()

    docx_export._ascii_table(
        doc,
        columns=[("欄位一", 10), ("欄位二", 10)],
        rows=[["會換行的很長內容", "value"]],
        indent="  ",
    )
    docx_export._render_health_pgx_section(doc, "藥物基因體學", {})

    assert "keepNext" not in doc.element.xml
